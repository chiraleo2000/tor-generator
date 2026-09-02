"""
Knowledge RAG — FastAPI Backend
Azure OpenAI + BM25 + PDF Extraction
"""

from fastapi import FastAPI, File, Form, UploadFile, HTTPException, Request
from fastapi.staticfiles import StaticFiles
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import pdfplumber, json, asyncio, httpx, re, math, uuid, time, os, shutil, hmac, io
from pathlib import Path
from typing import List, Optional
from dotenv import load_dotenv
from metadata_extractor import (
    build_source_metadata,
    empty_document_metadata,
    extract_document_metadata,
    flatten_for_kb,
    save_document_metadata,
    summarize_for_meta,
)

load_dotenv()  # โหลด .env จาก root ของโปรเจค

# ─────────────────────────────────────────────────────────────
BASE_DIR   = Path(__file__).parent
STATIC_DIR = BASE_DIR / "static"
KB_DIR     = Path(os.getenv("PAGEINDEX_KB_DIR", str(BASE_DIR / "knowledge"))).expanduser().resolve()
CONFIG_F   = BASE_DIR / "config.json"
SESSIONS_F = BASE_DIR / "sessions.json"
KB_DIR.mkdir(parents=True, exist_ok=True)


def _env_bool(name: str, default: bool = False) -> bool:
    raw = os.getenv(name)
    if raw is None:
        return default
    return raw.strip().lower() in {"1", "true", "yes", "on"}


PAGEINDEX_UI_ENABLED = _env_bool("PAGEINDEX_UI_ENABLED", False)

# ─── Session Memory ───────────────────────────────────────────
_sess_lock = asyncio.Lock()

def _load_sessions() -> dict:
    if SESSIONS_F.exists():
        try: return json.loads(SESSIONS_F.read_text(encoding="utf-8"))
        except: pass
    return {}

def _persist_sessions(sess: dict):
    # Keep last 100 sessions, 60 messages each
    if len(sess) > 100:
        for k in list(sess.keys())[:-100]: del sess[k]
    for k in sess:
        sess[k] = sess[k][-60:]
    SESSIONS_F.write_text(json.dumps(sess, ensure_ascii=False, indent=2), encoding="utf-8")

async def get_history(session_id: str, doc_id: str) -> list:
    sess = _load_sessions()
    return sess.get(f"{session_id}:{doc_id}", [])

async def push_message(session_id: str, doc_id: str, role: str, content: str):
    async with _sess_lock:
        sess = _load_sessions()
        key = f"{session_id}:{doc_id}"
        sess.setdefault(key, []).append({"role": role, "content": content})
        _persist_sessions(sess)

async def clear_history(session_id: str, doc_id: str):
    async with _sess_lock:
        sess = _load_sessions()
        key = f"{session_id}:{doc_id}"
        if key in sess: del sess[key]
        _persist_sessions(sess)

app = FastAPI(title="Knowledge RAG")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])


@app.middleware("http")
async def require_pageindex_api_key(request: Request, call_next):
    """Protect PageIndex APIs when PAGEINDEX_API_KEY is configured."""
    expected = os.getenv("PAGEINDEX_API_KEY", "").strip()
    if expected and request.url.path.startswith("/api/"):
        supplied = request.headers.get("X-API-Key", "")
        bearer = request.headers.get("Authorization", "")
        if not supplied and bearer.lower().startswith("bearer "):
            supplied = bearer[7:].strip()
        if not supplied or not hmac.compare_digest(supplied, expected):
            from fastapi.responses import JSONResponse
            return JSONResponse(status_code=401, content={"detail": "invalid PageIndex API key"})
    return await call_next(request)


class LinkUploadReq(BaseModel):
    url: str


class LineImportReq(BaseModel):
    message_id: Optional[str] = None
    message: Optional[dict] = None


class SearchReq(BaseModel):
    query: str
    doc_id: Optional[str] = None
    k: int = 6
    user_id: Optional[str] = None
    search_scope: str = "both"


class GraphSearchReq(BaseModel):
    query: str
    k: int = 8
    doc_limit: int = 6


# ─── Config ───────────────────────────────────────────────────
DEFAULT_CFG = {
    "endpoint":    "",
    "api_key":     "",
    "model":       "gpt-4o-mini",
    "api_version": "2024-08-01-preview",
}

def load_cfg() -> dict:
    # Priority: config.json (UI) → .env → default
    base = {
        "endpoint":           os.getenv("AZURE_OPENAI_ENDPOINT", ""),
        "api_key":            os.getenv("AZURE_OPENAI_API_KEY",  ""),
        "model":              os.getenv("AZURE_OPENAI_MODEL",    "gpt-4o-mini"),
        "api_version":        os.getenv("AZURE_OPENAI_API_VERSION", "2024-08-01-preview"),
        "doc_intel_endpoint": os.getenv("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT", ""),
        "doc_intel_key":      os.getenv("AZURE_DOCUMENT_INTELLIGENCE_KEY", ""),
    }
    if CONFIG_F.exists():
        saved = json.loads(CONFIG_F.read_text())
        for k, v in saved.items():
            if v: base[k] = v
    return base

def save_cfg(d: dict):
    current = load_cfg()
    current.update(d)
    CONFIG_F.write_text(json.dumps(current, indent=2))

# ─── Knowledge Base helpers ───────────────────────────────────
def load_kb(doc_id: str) -> dict | None:
    p = KB_DIR / doc_id / "knowledge_base.json"
    return json.loads(p.read_text(encoding="utf-8")) if p.exists() else None

def load_meta(doc_id: str) -> dict | None:
    p = KB_DIR / doc_id / "meta.json"
    return json.loads(p.read_text()) if p.exists() else None

def load_document_meta(doc_id: str) -> dict | None:
    p = KB_DIR / doc_id / "document_meta.json"
    return json.loads(p.read_text(encoding="utf-8")) if p.exists() else None

def load_graph(doc_id: str) -> dict | None:
    p = KB_DIR / doc_id / "graph.json"
    return json.loads(p.read_text(encoding="utf-8")) if p.exists() else None

def save_meta(doc_id: str, meta: dict):
    (KB_DIR / doc_id / "meta.json").write_text(json.dumps(meta, ensure_ascii=False))

def save_initial_document_metadata(doc_id: str, source: dict):
    save_document_metadata(KB_DIR / doc_id, empty_document_metadata(source))

async def enrich_and_save_document_metadata(
    doc_id: str,
    text: str,
    cfg: dict,
    *,
    source: dict | None = None,
    fallback_title: str | None = None,
    kb: dict | None = None,
) -> dict:
    doc_dir = KB_DIR / doc_id
    existing = load_document_meta(doc_id) or {}
    source = source or existing.get("source") or build_source_metadata(
        doc_id=doc_id,
        source_kind=(load_meta(doc_id) or {}).get("source_kind") or "pdf",
        source_origin=(load_meta(doc_id) or {}).get("source_origin") or "manual_upload",
    )
    metadata = await extract_document_metadata(text, cfg, source, fallback_title=fallback_title)
    save_document_metadata(doc_dir, metadata)

    flat = flatten_for_kb(metadata)
    if kb is not None:
        kb.setdefault("meta", {})
        kb["meta"]["doc_meta"] = flat
        if flat.get("document_title"):
            kb["meta"]["doc_title"] = flat["document_title"]

    meta = load_meta(doc_id) or {"doc_id": doc_id}
    meta.update({k: v for k, v in summarize_for_meta(metadata).items() if v is not None})
    if flat.get("document_title"):
        meta["filename"] = flat["document_title"]
    save_meta(doc_id, meta)
    return metadata

def all_docs() -> list:
    docs = []
    for d in sorted(KB_DIR.iterdir()):
        if d.is_dir():
            m = load_meta(d.name)
            if m:
                docs.append(m)
    return docs


LINE_HUB_API_BASE = os.getenv(
    "LINE_HUB_API_BASE",
    "https://rag-line-hub.s65122251007.workers.dev/api/v1",
).rstrip("/")


def find_doc_by_external_id(external_id: str) -> dict | None:
    if not external_id:
        return None
    for meta in all_docs():
        if meta.get("external_id") == external_id:
            return meta
    return None


def _extract_line_url(item: dict) -> str | None:
    msg = item.get("message") or item
    if msg.get("url"):
        return msg["url"]
    links = msg.get("links") or item.get("links") or []
    if links:
        return links[0]
    content = msg.get("content") or item.get("content") or msg.get("markdown") or item.get("markdown") or ""
    m = re.search(r"https?://\S+", content)
    return m.group(0).rstrip(").,]\"'") if m else None


def _line_message_label(item: dict, url: str) -> str:
    sender = (item.get("sender") or {}).get("name") or "LINE"
    received = item.get("receivedAt") or item.get("received_at") or ""
    host = re.sub(r"^https?://", "", url).split("/")[0]
    return f"{sender} · {host} · {received[:10]}".strip(" ·")[:160]


def create_line_link_doc(item: dict, source: str = "pull") -> dict:
    external_id = item.get("id") or item.get("external_id") or (item.get("message") or {}).get("id")
    existing = find_doc_by_external_id(external_id)
    if existing:
        return {"doc_id": existing["doc_id"], "filename": existing.get("filename"), "existing": True}

    url = _extract_line_url(item)
    if not url:
        raise ValueError("LINE message นี้ไม่มี URL สำหรับนำเข้า")

    doc_id = uuid.uuid4().hex[:8]
    doc_dir = KB_DIR / doc_id
    doc_dir.mkdir()
    (doc_dir / "source.url").write_text(url, encoding="utf-8")
    (doc_dir / "line_message.json").write_text(json.dumps(item, ensure_ascii=False, indent=2), encoding="utf-8")

    meta = {
        "doc_id": doc_id,
        "filename": _line_message_label(item, url),
        "source_kind": "link",
        "source_origin": "line_api",
        "source_url": url,
        "external_id": external_id,
        "line_import_source": source,
        "line_sender": item.get("sender"),
        "line_received_at": item.get("receivedAt") or item.get("received_at"),
        "line_status": item.get("status") or "pending",
        "status": "uploaded",
        "created": time.time(),
        "total_records": 0,
    }
    save_meta(doc_id, meta)
    source = build_source_metadata(
        doc_id=doc_id,
        source_kind="link",
        source_origin="line_api",
        source_url=url,
        original_filename=meta["filename"],
        extra={
            "external_id": external_id,
            "line_received_at": meta.get("line_received_at"),
        },
    )
    save_initial_document_metadata(doc_id, source)
    return {"doc_id": doc_id, "filename": meta["filename"], "source_url": url, "existing": False}


async def fetch_line_message(message_id: str) -> dict:
    async with httpx.AsyncClient(timeout=20) as client:
        resp = await client.get(f"{LINE_HUB_API_BASE}/messages/{message_id}")
        resp.raise_for_status()
        return resp.json().get("data") or {}


async def ack_line_message(message_id: str) -> dict | None:
    if not message_id:
        return None
    async with httpx.AsyncClient(timeout=20) as client:
        resp = await client.post(f"{LINE_HUB_API_BASE}/messages/{message_id}/ack")
        resp.raise_for_status()
        return resp.json()

# ─── BM25 ────────────────────────────────────────────────────
K1, BV = 1.5, 0.75

def is_thai(c): return "฀" <= c <= "๿"

def tokenize(text: str) -> list[str]:
    if not text: return []
    words = [t for t in re.split(r'[\s;,.:()[\]/|·•–\-–—"\']+', text.lower()) if len(t) > 1]
    out = list(words)
    for w in words:
        if len(w) >= 3 and any(is_thai(c) for c in w):
            out += [w[i:i+3] for i in range(len(w) - 2)]
    return out

class BM25:
    def __init__(self, records: list):
        self.records = records
        self.corpus  = [tokenize(r.get("search_text") or " ".join(filter(None, [
            r.get("section_id",""), r.get("title",""), r.get("summary",""),
            r.get("details",""), " ".join(r.get("keywords",[]))
        ]))) for r in records]
        N = len(self.corpus)
        self.avgdl = sum(len(d) for d in self.corpus) / max(N, 1)
        df: dict = {}
        for d in self.corpus:
            for t in set(d): df[t] = df.get(t, 0) + 1
        self.idf = {t: math.log((N - n + 0.5) / (n + 0.5) + 1) for t, n in df.items()}

    def search(self, query: str, k: int = 6) -> list[tuple]:
        qt = tokenize(query)
        scored = []
        for i, doc in enumerate(self.corpus):
            dl = len(doc); freq: dict = {}
            for t in doc: freq[t] = freq.get(t, 0) + 1
            s = sum(
                self.idf.get(t, 0) * (freq.get(t, 0) * (K1 + 1))
                / (freq.get(t, 0) + K1 * (1 - BV + BV * dl / self.avgdl))
                for t in qt if freq.get(t, 0)
            ) * (self.records[i].get("boost") or 1.0)
            if s > 0: scored.append((s, i))
        scored.sort(reverse=True)
        return [(self.records[i], round(s, 2)) for s, i in scored[:k]]

_bm25_cache: dict[str, BM25] = {}
_graph_tasks: dict[str, asyncio.Task] = {}

def get_bm25(doc_id: str) -> BM25 | None:
    if doc_id not in _bm25_cache:
        kb = load_kb(doc_id)
        if kb: _bm25_cache[doc_id] = BM25(kb["records"])
    return _bm25_cache.get(doc_id)

def _graph_running(doc_id: str) -> bool:
    task = _graph_tasks.get(doc_id)
    return bool(task and not task.done())

def _set_graph_meta(doc_id: str, **updates):
    meta = load_meta(doc_id) or {"doc_id": doc_id}
    meta.update(updates)
    save_meta(doc_id, meta)

async def _build_graph_job(doc_id: str, cfg: dict):
    try:
        kb = load_kb(doc_id)
        if not kb:
            raise RuntimeError("Knowledge base not ready")
        from graph_extractor import build_relation_graph

        graph = await asyncio.to_thread(build_relation_graph, kb, cfg, doc_id)
        (KB_DIR / doc_id / "graph.json").write_text(
            json.dumps(graph, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        _set_graph_meta(
            doc_id,
            graph_ready=True,
            graph_status="ready",
            graph_error=None,
            graph_nodes=len(graph.get("nodes", [])),
            graph_edges=len(graph.get("edges", [])),
            graph_updated=time.time(),
        )
    except Exception as exc:
        _set_graph_meta(
            doc_id,
            graph_ready=False,
            graph_status="error",
            graph_error=str(exc),
            graph_updated=time.time(),
        )
        print(f"[graph] build failed for {doc_id}: {exc}")
    finally:
        task = asyncio.current_task()
        if task and _graph_tasks.get(doc_id) is task:
            _graph_tasks.pop(doc_id, None)

def schedule_graph_build(doc_id: str, cfg: dict, force: bool = False) -> bool:
    if _graph_running(doc_id):
        return False
    graph_path = KB_DIR / doc_id / "graph.json"
    if graph_path.exists() and not force:
        return False
    _set_graph_meta(
        doc_id,
        graph_ready=False,
        graph_status="processing",
        graph_error=None,
        graph_started=time.time(),
    )
    _graph_tasks[doc_id] = asyncio.create_task(_build_graph_job(doc_id, dict(cfg)))
    return True


# ─── Global graph + graph-first retrieval ─────────────────────
def _doc_title(meta: dict, doc_id: str) -> str:
    return meta.get("filename") or meta.get("document_title") or doc_id


def _norm_graph_name(value: str) -> str:
    value = re.sub(r"\s+", " ", str(value or "").strip().lower())
    value = re.sub(r"^[\W_]+|[\W_]+$", "", value, flags=re.UNICODE)
    return value or "unknown"


def _append_unique(items: list, value):
    if value and value not in items:
        items.append(value)


def _doc_node(doc_id: str, meta: dict) -> dict:
    title = _doc_title(meta, doc_id)
    desc_parts = [
        meta.get("summary"),
        meta.get("agency"),
        meta.get("ministry"),
        meta.get("collection"),
        meta.get("source_type"),
    ]
    return {
        "id": f"doc:{doc_id}",
        "name": title,
        "type": "document",
        "group": "Document",
        "description": " · ".join(str(x) for x in desc_parts if x),
        "size": 2,
        "doc_id": doc_id,
        "doc_ids": [doc_id],
        "docs": [{
            "doc_id": doc_id,
            "title": title,
            "source_kind": meta.get("source_kind"),
            "source_url": meta.get("source_url"),
        }],
    }


def _skip_global_entity(raw: dict) -> bool:
    technical = {
        "pdf object",
        "document structure",
        "image",
        "image format",
        "encoding",
        "filter",
        "font",
        "color space",
    }
    node_type = _norm_graph_name(raw.get("type") or raw.get("category") or "")
    group = _norm_graph_name(raw.get("group") or node_type)
    name = _norm_graph_name(raw.get("name") or raw.get("id") or raw.get("label") or "")
    if node_type in technical or group in technical:
        return True
    return bool(re.search(r"\b(pdf object|xobject|flatedecode|jfif|devicegray|winansi)\b", name))


def _build_relation_global_graph() -> dict:
    nodes: dict[str, dict] = {}
    edges: dict[str, dict] = {}
    entity_docs: dict[str, set] = {}
    doc_entities: dict[str, set] = {}
    docs_used = []

    def add_edge(source: str, target: str, label: str, doc_id: str | None = None,
                 description: str = "", weight: float = 1.0):
        if not source or not target or source == target:
            return
        edge_id = f"{source}|{label}|{target}"
        if edge_id not in edges:
            edges[edge_id] = {
                "id": edge_id,
                "source": source,
                "target": target,
                "label": label,
                "type": label,
                "description": description,
                "weight": 0.0,
                "doc_ids": [],
            }
        edge = edges[edge_id]
        edge["weight"] = round(float(edge.get("weight") or 0) + float(weight or 1), 2)
        if description and description not in edge.get("description", ""):
            edge["description"] = (edge.get("description") + " / " + description).strip(" /")[:500]
        if doc_id:
            _append_unique(edge["doc_ids"], doc_id)

    for meta in all_docs():
        doc_id = meta.get("doc_id")
        if not doc_id or meta.get("status") != "ready":
            continue
        doc_key = f"doc:{doc_id}"
        nodes[doc_key] = _doc_node(doc_id, meta)
        doc_entities.setdefault(doc_id, set())
        docs_used.append({
            "doc_id": doc_id,
            "title": _doc_title(meta, doc_id),
            "source_kind": meta.get("source_kind"),
            "source_url": meta.get("source_url"),
        })

        graph = load_graph(doc_id) or {}
        local_to_global: dict[str, str] = {}
        for raw in graph.get("nodes") or []:
            if _skip_global_entity(raw):
                continue
            name = str(raw.get("name") or raw.get("id") or "").strip()
            if not name:
                continue
            entity_key = f"entity:{_norm_graph_name(name)}"
            local_to_global[str(raw.get("id") or name)] = entity_key
            entity_docs.setdefault(entity_key, set()).add(doc_id)
            doc_entities[doc_id].add(entity_key)
            if entity_key not in nodes:
                nodes[entity_key] = {
                    "id": entity_key,
                    "name": name,
                    "type": raw.get("type") or "Entity",
                    "group": raw.get("group") or raw.get("type") or "Entity",
                    "description": raw.get("description") or "",
                    "size": 1,
                    "doc_ids": [],
                    "docs": [],
                    "mention_count": 0,
                }
            node = nodes[entity_key]
            node["mention_count"] = int(node.get("mention_count") or 0) + 1
            node["size"] = max(1, min(8, len(entity_docs[entity_key]) + 1))
            _append_unique(node["doc_ids"], doc_id)
            _append_unique(node["docs"], nodes[doc_key]["docs"][0])
            if not node.get("description") and raw.get("description"):
                node["description"] = raw.get("description")
            add_edge(doc_key, entity_key, "mentions", doc_id, weight=0.35)

        for raw in graph.get("edges") or []:
            source = local_to_global.get(str(raw.get("source") or ""), f"entity:{_norm_graph_name(raw.get('source') or '')}")
            target = local_to_global.get(str(raw.get("target") or ""), f"entity:{_norm_graph_name(raw.get('target') or '')}")
            if source in nodes and target in nodes:
                add_edge(
                    source,
                    target,
                    str(raw.get("label") or raw.get("type") or "related_to"),
                    doc_id,
                    str(raw.get("description") or ""),
                    float(raw.get("weight") or 1),
                )

    doc_ids = list(doc_entities)
    for i, a in enumerate(doc_ids):
        for b in doc_ids[i + 1:]:
            shared = doc_entities[a] & doc_entities[b]
            if len(shared) < 2:
                continue
            label = f"shares {len(shared)} entities"
            add_edge(f"doc:{a}", f"doc:{b}", label, None, weight=min(4, len(shared) / 2))
            edge = edges[f"doc:{a}|{label}|doc:{b}"]
            edge["doc_ids"] = [a, b]

    groups = {}
    for node in nodes.values():
        groups.setdefault(node.get("group") or "Entity", 0)
        groups[node.get("group") or "Entity"] += 1

    return {
        "scope": "global",
        "doc_id": None,
        "generated_at": time.time(),
        "engine": "global-pageindex-graph",
        "docs": docs_used,
        "nodes": list(nodes.values()),
        "edges": list(edges.values()),
        "groups": [{"id": k, "label": k, "count": v} for k, v in sorted(groups.items())],
    }


def _score_text(query_tokens: list[str], text: str) -> float:
    tokens = tokenize(text)
    if not tokens:
        return 0.0
    freq = {}
    for token in tokens:
        freq[token] = freq.get(token, 0) + 1
    score = sum(freq.get(token, 0) for token in set(query_tokens))
    query = " ".join(query_tokens).strip().lower()
    if query and query in str(text or "").lower():
        score += 3
    return float(score)


def _record_concepts(rec: dict) -> list[str]:
    concepts = []
    meta = rec.get("section_meta") or {}
    for value in rec.get("keywords") or []:
        _append_unique(concepts, str(value).strip())
    for value in meta.get("topic") or []:
        _append_unique(concepts, str(value).strip())
    for value in meta.get("use_case") or []:
        _append_unique(concepts, str(value).strip())
    for value in (meta.get("legal_domain"), rec.get("chapter")):
        if value:
            _append_unique(concepts, str(value).strip())
    return [c for c in concepts if len(c) >= 2][:10]


def build_global_graph() -> dict:
    """Build a cross-document graph from PageIndex records.

    This is intentionally separate from per-document Hyper-Extract graphs:
    the global graph maps documents -> PageIndex sections -> shared concepts,
    then adds cross-document similarity edges from shared PageIndex concepts.
    """
    nodes: dict[str, dict] = {}
    edges: dict[str, dict] = {}
    docs_used = []
    section_keys: dict[tuple[str, str], str] = {}
    section_concepts: dict[str, set] = {}
    doc_concepts: dict[str, set] = {}

    def add_edge(source: str, target: str, label: str, weight: float = 1.0, **extra):
        if not source or not target or source == target:
            return
        edge_id = f"{source}|{label}|{target}"
        if edge_id not in edges:
            edges[edge_id] = {
                "id": edge_id,
                "source": source,
                "target": target,
                "label": label,
                "type": label,
                "description": extra.get("description", ""),
                "weight": 0.0,
                "doc_ids": [],
            }
        edge = edges[edge_id]
        edge["weight"] = round(float(edge.get("weight") or 0) + float(weight or 1), 2)
        for doc_id in extra.get("doc_ids") or []:
            _append_unique(edge["doc_ids"], doc_id)

    for meta in all_docs():
        doc_id = meta.get("doc_id")
        if not doc_id or meta.get("status") != "ready":
            continue
        kb = load_kb(doc_id)
        if not kb:
            continue

        doc_info = {
            "doc_id": doc_id,
            "title": _doc_title(meta, doc_id),
            "source_kind": meta.get("source_kind"),
            "source_url": meta.get("source_url"),
        }
        docs_used.append(doc_info)
        doc_key = f"doc:{doc_id}"
        nodes[doc_key] = _doc_node(doc_id, meta)
        doc_concepts.setdefault(doc_id, set())

        for rec in kb.get("records") or []:
            sid = str(rec.get("section_id") or "").strip()
            if not sid:
                continue
            sec_key = f"section:{doc_id}:{sid}"
            section_keys[(doc_id, sid)] = sec_key
            title = rec.get("title") or sid
            summary = rec.get("summary") or rec.get("details") or ""
            nodes[sec_key] = {
                "id": sec_key,
                "name": f"[{sid}] {title}",
                "type": "section",
                "group": "PageIndex Section",
                "description": str(summary)[:450],
                "size": 1,
                "doc_id": doc_id,
                "doc_ids": [doc_id],
                "section_id": sid,
                "page_start": rec.get("page_start"),
                "page_end": rec.get("page_end"),
                "docs": [doc_info],
            }
            add_edge(doc_key, sec_key, "has_section", 0.45, doc_ids=[doc_id])

            parent_sid = ".".join(sid.split(".")[:-1])
            if parent_sid and (doc_id, parent_sid) in section_keys:
                add_edge(section_keys[(doc_id, parent_sid)], sec_key, "has_child", 0.7, doc_ids=[doc_id])

            concepts = set()
            for concept in _record_concepts(rec):
                concept_key = f"concept:{_norm_graph_name(concept)}"
                concepts.add(concept_key)
                doc_concepts[doc_id].add(concept_key)
                if concept_key not in nodes:
                    nodes[concept_key] = {
                        "id": concept_key,
                        "name": concept,
                        "type": "concept",
                        "group": "Concept",
                        "description": "PageIndex concept/keyword shared by sections",
                        "size": 1,
                        "doc_ids": [],
                        "docs": [],
                        "section_refs": [],
                    }
                concept_node = nodes[concept_key]
                _append_unique(concept_node["doc_ids"], doc_id)
                _append_unique(concept_node["docs"], doc_info)
                _append_unique(concept_node["section_refs"], {"doc_id": doc_id, "section_id": sid, "title": title})
                concept_node["size"] = max(1, min(8, len(concept_node["doc_ids"]) + 1))
                add_edge(sec_key, concept_key, "mentions", 0.7, doc_ids=[doc_id])
            section_concepts[sec_key] = concepts

    doc_ids = list(doc_concepts)
    for i, a in enumerate(doc_ids):
        for b in doc_ids[i + 1:]:
            shared = doc_concepts[a] & doc_concepts[b]
            if len(shared) >= 3:
                add_edge(f"doc:{a}", f"doc:{b}", f"shares {len(shared)} pageindex concepts", min(5, len(shared) / 2), doc_ids=[a, b])

    similar_sections = []
    section_items = list(section_concepts.items())
    for i, (a_key, a_concepts) in enumerate(section_items):
        a_doc = nodes[a_key].get("doc_id")
        for b_key, b_concepts in section_items[i + 1:]:
            if a_doc == nodes[b_key].get("doc_id"):
                continue
            shared = a_concepts & b_concepts
            if len(shared) >= 2:
                similar_sections.append((len(shared), a_key, b_key))
    for shared_count, a_key, b_key in sorted(similar_sections, reverse=True)[:220]:
        add_edge(
            a_key,
            b_key,
            f"similar_by_{shared_count}_concepts",
            min(4, shared_count),
            doc_ids=[nodes[a_key].get("doc_id"), nodes[b_key].get("doc_id")],
        )

    groups = {}
    for node in nodes.values():
        groups.setdefault(node.get("group") or "Entity", 0)
        groups[node.get("group") or "Entity"] += 1

    return {
        "scope": "global_pageindex",
        "doc_id": None,
        "generated_at": time.time(),
        "engine": "pageindex-graph",
        "docs": docs_used,
        "nodes": list(nodes.values()),
        "edges": list(edges.values()),
        "groups": [{"id": k, "label": k, "count": v} for k, v in sorted(groups.items())],
    }


def graph_first_search(
    query: str,
    k: int = 8,
    doc_limit: int = 6,
    allowed_doc_ids: set[str] | None = None,
) -> dict:
    q = (query or "").strip()
    q_tokens = tokenize(q)
    if not q_tokens:
        return {"query": q, "hits": [], "graph_matches": []}

    graph = build_global_graph()
    node_by_id = {n["id"]: n for n in graph.get("nodes", [])}
    doc_scores: dict[str, float] = {}
    matches = []

    for node in graph.get("nodes", []):
        node_text = " ".join(filter(None, [
            node.get("name"),
            node.get("type"),
            node.get("group"),
            node.get("description"),
            " ".join(d.get("title", "") for d in node.get("docs") or []),
        ]))
        score = _score_text(q_tokens, node_text)
        if score <= 0:
            continue
        if node.get("type") == "document" and node.get("doc_id"):
            doc_scores[node["doc_id"]] = doc_scores.get(node["doc_id"], 0) + score * 2.5
        for doc_id in node.get("doc_ids") or []:
            doc_scores[doc_id] = doc_scores.get(doc_id, 0) + score * 1.5
        matches.append({
            "kind": "node",
            "id": node.get("id"),
            "label": node.get("name"),
            "group": node.get("group"),
            "score": round(score, 2),
            "doc_ids": node.get("doc_ids") or [],
        })

    for edge in graph.get("edges", []):
        source = node_by_id.get(edge.get("source"), {})
        target = node_by_id.get(edge.get("target"), {})
        edge_text = " ".join(filter(None, [
            source.get("name"),
            edge.get("label"),
            target.get("name"),
            edge.get("description"),
        ]))
        score = _score_text(q_tokens, edge_text)
        if score <= 0:
            continue
        for doc_id in edge.get("doc_ids") or []:
            doc_scores[doc_id] = doc_scores.get(doc_id, 0) + score
        matches.append({
            "kind": "edge",
            "id": edge.get("id"),
            "label": edge.get("label"),
            "source": source.get("name") or edge.get("source"),
            "target": target.get("name") or edge.get("target"),
            "score": round(score, 2),
            "doc_ids": edge.get("doc_ids") or [],
        })

    ready_docs = [
        d["doc_id"]
        for d in all_docs()
        if d.get("status") == "ready"
        and (allowed_doc_ids is None or d["doc_id"] in allowed_doc_ids)
    ]
    if allowed_doc_ids is not None:
        visible_matches = []
        for match in matches:
            visible_doc_ids = [
                doc_id
                for doc_id in match.get("doc_ids") or []
                if doc_id in allowed_doc_ids
            ]
            if not visible_doc_ids:
                continue
            visible_matches.append({**match, "doc_ids": visible_doc_ids})
        matches = visible_matches
    ranked_docs = [
        doc_id for doc_id, _ in sorted(doc_scores.items(), key=lambda item: item[1], reverse=True)
        if doc_id in ready_docs
    ][: max(1, min(doc_limit, 20))]
    if not ranked_docs:
        ranked_docs = ready_docs

    matches.sort(key=lambda item: item["score"], reverse=True)
    hits = []
    for doc_id in ranked_docs:
        bm25 = get_bm25(doc_id)
        if not bm25:
            continue
        meta = load_meta(doc_id) or {}
        graph_boost = min(doc_scores.get(doc_id, 0), 8)
        for rec, score in bm25.search(q, k=k):
            hits.append({
                "doc_id": doc_id,
                "doc_title": _doc_title(meta, doc_id),
                "source_kind": meta.get("source_kind"),
                "source_origin": meta.get("source_origin"),
                "source_url": meta.get("source_url"),
                "section_id": rec.get("section_id"),
                "title": rec.get("title"),
                "summary": rec.get("summary"),
                "details": rec.get("details", "")[:1200],
                "keywords": rec.get("keywords") or [],
                "score": round(float(score) + graph_boost, 2),
                "graph_boost": round(graph_boost, 2),
            })

    hits.sort(key=lambda item: item["score"], reverse=True)
    return {
        "query": q,
        "expanded_query": q,
        "graph_matches": matches[:10],
        "candidate_doc_ids": ranked_docs,
        "hits": hits[: max(1, min(k, 30))],
    }


def prune_graph_for_ui(graph: dict, max_sections: int = 80, max_concepts: int = 50) -> dict:
    nodes = graph.get("nodes") or []
    edges = graph.get("edges") or []
    degree = {}
    for edge in edges:
        degree[edge.get("source")] = degree.get(edge.get("source"), 0) + 1
        degree[edge.get("target")] = degree.get(edge.get("target"), 0) + 1

    docs = [n for n in nodes if n.get("type") == "document"]
    sections = [n for n in nodes if n.get("type") == "section"]
    concepts = [n for n in nodes if n.get("type") == "concept"]

    sections = sorted(
        sections,
        key=lambda n: (len(n.get("doc_ids") or []), degree.get(n.get("id"), 0), len(str(n.get("description") or ""))),
        reverse=True,
    )[:max_sections]
    concepts = sorted(
        concepts,
        key=lambda n: (len(n.get("doc_ids") or []), degree.get(n.get("id"), 0), len(n.get("section_refs") or [])),
        reverse=True,
    )[:max_concepts]

    keep = {n["id"] for n in docs + sections + concepts}
    pruned_edges = [e for e in edges if e.get("source") in keep and e.get("target") in keep]
    pruned_nodes = []
    for node in nodes:
        if node.get("id") not in keep:
            continue
        compact = dict(node)
        if len(compact.get("docs") or []) > 8:
            compact["docs"] = compact["docs"][:8]
            compact["docs_truncated"] = True
        if len(compact.get("section_refs") or []) > 10:
            compact["section_refs"] = compact["section_refs"][:10]
            compact["section_refs_truncated"] = True
        if compact.get("description"):
            compact["description"] = str(compact["description"])[:450]
        pruned_nodes.append(compact)

    groups = {}
    for node in pruned_nodes:
        groups.setdefault(node.get("group") or "Entity", 0)
        groups[node.get("group") or "Entity"] += 1

    out = dict(graph)
    out["nodes"] = pruned_nodes
    out["edges"] = pruned_edges
    out["groups"] = [{"id": k, "label": k, "count": v} for k, v in sorted(groups.items())]
    out["truncated"] = len(pruned_nodes) < len(nodes) or len(pruned_edges) < len(edges)
    out["full_counts"] = {"nodes": len(nodes), "edges": len(edges)}
    return out

# ─── Azure OpenAI ─────────────────────────────────────────────
async def azure_stream(messages: list, cfg: dict):
    """Yield raw SSE data lines from Azure OpenAI streaming."""
    ep  = cfg["endpoint"].rstrip("/")
    url = f"{ep}/openai/deployments/{cfg['model']}/chat/completions?api-version={cfg['api_version']}"
    hdrs = {"api-key": cfg["api_key"], "Content-Type": "application/json"}
    body = {"messages": messages, "stream": True, "max_completion_tokens": 1800, "temperature": 0.1}

    async with httpx.AsyncClient(timeout=90) as c:
        async with c.stream("POST", url, json=body, headers=hdrs) as r:
            r.raise_for_status()
            async for line in r.aiter_lines():
                if line.startswith("data: ") and line != "data: [DONE]":
                    yield line[6:]

async def azure_once(messages: list, cfg: dict) -> str:
    """Single (non-streaming) call → return full text."""
    ep  = cfg["endpoint"].rstrip("/")
    url = f"{ep}/openai/deployments/{cfg['model']}/chat/completions?api-version={cfg['api_version']}"
    hdrs = {"api-key": cfg["api_key"], "Content-Type": "application/json"}
    body = {"messages": messages, "max_completion_tokens": 800, "temperature": 0.1}

    async with httpx.AsyncClient(timeout=60) as c:
        r = await c.post(url, json=body, headers=hdrs)
        r.raise_for_status()
        return r.json()["choices"][0]["message"]["content"]

# ─── PDF Extraction helpers ───────────────────────────────────
SEC_RE = re.compile(r"^(\d+(?:\.\d+){0,4})\s{1,4}(.{3,80})$")

async def extract_pages(path: str, cfg: dict) -> list[dict]:
    """Read PDF pages via Azure Document Intelligence (or local fallback)."""
    from ocr import read_pdf_pages
    texts = await read_pdf_pages(path, cfg)
    return [{"page": i + 1, "text": t} for i, t in enumerate(texts)]

def parse_sections(pages: list[dict]) -> list[dict]:
    sections, cur, cur_lines = [], None, []

    def flush():
        if cur:
            cur["content"] = "\n".join(cur_lines).strip()
            sections.append(dict(cur))

    for pg in pages:
        for line in pg["text"].split("\n"):
            line = line.strip()
            if not line: continue
            m = SEC_RE.match(line)
            if m and len(m.group(1)) <= 12:
                flush()
                cur = {"section_id": m.group(1), "title": m.group(2).strip(), "page": pg["page"]}
                cur_lines = []
            elif cur:
                cur_lines.append(line)

    flush()
    return sections

# ─── RAG กลาง Extraction Prompts ──────────────────────────────
CLASSIFY_PROMPT = """วิเคราะห์เนื้อหาเอกสารนี้และระบุ metadata ระดับเอกสาร ตอบเป็น JSON object เดียว ไม่มีข้อความอื่น

ค่าที่ใช้ได้:
- collection: Strategy | Agency | News | Bidding | Procurement | Vendor | Political | Relationship | Compliance | Legal | Risk
- source_type: Official Plan | Law | Regulation | Circular | News | TOR | Award Notice | Court Case | Internal Analysis
- government_tier: ส่วนกลาง | ภูมิภาค | ท้องถิ่น
- legal_domain: Procurement | PDPA | Cybersecurity | Cloud | Digital Government | Contract | Labor | Finance (หรือ null)
- validity: Current | Historical | Superseded
- reliability_level: Official | Verified | News | Secondary | Internal Analysis
- confidentiality: Public | Internal | Restricted
- use_case: Engagement | Presale | Compliance Screening | Market Intelligence
- project_type: e-Service | Data Platform | Call Center | OSS | e-Document | AI Assistant | อื่นๆ | null
- relationship_type: Awarded Contract | MOU | Seminar | Advisor | Past Project | News Mention
- risk_flags: locked_spec | appeal | complaint | legal_dispute | delivery_risk | privacy_risk | cybersecurity_risk
- procurement_status: TOR | announced | awarded | cancelled | appeal | contract_signed | null
- fact_vs_analysis: Fact | Analysis | Mixed

ตัวอย่าง JSON ที่ต้องตอบ (ห้ามมีข้อความอื่น):
{"document_title":"","collection":"Compliance","source_type":"Law","government_tier":"ส่วนกลาง","ministry":null,"agency":null,"related_agencies":[],"legal_domain":"Procurement","topic":[],"project_type":null,"vendor_si":[],"relationship_type":[],"risk_flags":[],"political_context":null,"fact_vs_analysis":"Fact","validity":"Current","reliability_level":"Official","confidentiality":"Public","use_case":["Compliance Screening"],"published_date":null,"last_reviewed":null,"geography":null,"procurement_status":null}"""

_DOC_META_DEFAULTS = {
    "collection": "Compliance", "source_type": "Official Plan",
    "government_tier": "ส่วนกลาง", "ministry": None, "agency": None,
    "related_agencies": [], "legal_domain": None, "topic": [],
    "project_type": None, "vendor_si": [], "relationship_type": [],
    "risk_flags": [], "political_context": None, "fact_vs_analysis": "Fact", "validity": "Current",
    "reliability_level": "Official", "confidentiality": "Public",
    "use_case": [], "published_date": None, "last_reviewed": None,
    "geography": None, "procurement_status": None,
}

async def classify_document(pages: list[dict], cfg: dict) -> dict:
    """Classify document once → return document-level metadata dict."""
    sample = "\n".join(p["text"] for p in pages[:5])[:3500]
    prompt = f"{CLASSIFY_PROMPT}\n\nเนื้อหาเอกสาร:\n{sample}"
    try:
        raw = await azure_once([
            {"role": "system", "content": "คุณเป็นผู้เชี่ยวชาญจำแนกเอกสารภาครัฐไทย ตอบ JSON เท่านั้น"},
            {"role": "user",   "content": prompt}
        ], cfg)
        raw = raw.strip()
        if raw.startswith("```"):
            raw = re.sub(r"^```[a-z]*\n?", "", raw); raw = re.sub(r"\n?```$", "", raw)
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if m:
            result = json.loads(m.group())
            merged = dict(_DOC_META_DEFAULTS)
            merged.update({k: v for k, v in result.items() if v is not None})
            return merged
    except Exception:
        pass
    return dict(_DOC_META_DEFAULTS)

ENRICH_SYS = "คุณเป็นผู้เชี่ยวชาญด้านกฎหมาย ระเบียบ และการจัดซื้อจัดจ้างภาครัฐไทย ตอบ JSON เท่านั้น"

async def enrich_batch(batch: list[dict], cfg: dict) -> list[dict]:
    """Enrich sections: summary, details, keywords + section-level RAG metadata."""
    n = len(batch)
    items = "\n---\n".join(
        f"[{i+1}] section_id={s['section_id']} title={s['title']}\n{s['content'][:500]}"
        for i, s in enumerate(batch)
    )
    prompt = (
        f"วิเคราะห์ {n} sections ต่อไปนี้ ตอบเป็น JSON array ขนาด {n} elements เท่านั้น (ไม่มีข้อความอื่น)\n\n"
        f"{items}\n\n"
        f"แต่ละ element มี field ดังนี้ (ตัวอย่าง):\n"
        '{"summary":"สรุป 2-3 ประโยคภาษาไทย","details":"รายละเอียดสำคัญ ≤600 ตัวอักษร","keywords":["คำ1","คำ2"],'
        '"legal_domain":null,"topic":["หัวข้อ"],"use_case":["Compliance Screening"],"boost":1.0}\n\n'
        "legal_domain ค่าที่ใช้ได้: Procurement|PDPA|Cybersecurity|Cloud|Digital Government|Contract|Labor|Finance หรือ null\n"
        "use_case ค่าที่ใช้ได้: Engagement|Presale|Compliance Screening|Market Intelligence\n"
        "boost: 2.0=บทนิยาม/หลักการหลัก, 1.5=วิธีการ/ข้อกำหนด, 1.0=ทั่วไป"
    )
    try:
        raw = await azure_once([
            {"role": "system", "content": ENRICH_SYS},
            {"role": "user",   "content": prompt}
        ], cfg)
        raw = raw.strip()
        if raw.startswith("```"):
            raw = re.sub(r"^```[a-z]*\n?", "", raw); raw = re.sub(r"\n?```$", "", raw)
        m = re.search(r"\[.*\]", raw, re.DOTALL)
        if m:
            data = json.loads(m.group())
            if len(data) == n: return data
            # Pad if LLM returned fewer items
            while len(data) < n:
                data.append({"summary": batch[len(data)]["title"], "details": "", "keywords": [],
                             "legal_domain": None, "topic": [], "use_case": [], "boost": 1.0})
            return data[:n]
    except Exception:
        pass
    return [{"summary": s["title"], "details": "", "keywords": [],
             "legal_domain": None, "topic": [], "use_case": [], "boost": 1.0}
            for s in batch]


def _fallback_example_queries(records: list[dict], doc_title: str = "") -> list[str]:
    titles = [r.get("title", "").strip() for r in records if r.get("title")]
    keywords = []
    for r in records:
        for kw in r.get("keywords") or []:
            if kw and kw not in keywords:
                keywords.append(kw)

    queries = []
    if doc_title:
        queries.append(f"สรุปประเด็นสำคัญของ {doc_title} คืออะไร?")
    for seed in (titles[:4] + keywords[:4]):
        seed = re.sub(r"\s+", " ", str(seed)).strip()[:80]
        if seed:
            queries.append(f"เอกสารนี้กล่าวถึง {seed} อย่างไร?")
    queries.extend([
        "มีข้อกำหนดหรือเงื่อนไขสำคัญอะไรที่ควรรู้?",
        "มีหน่วยงานหรือบุคคลใดเกี่ยวข้องบ้าง?",
        "ประเด็นใดควรนำไปใช้ติดตามต่อ?",
    ])

    deduped = []
    for q in queries:
        if q not in deduped:
            deduped.append(q)
    return deduped[:8]


async def generate_example_queries(kb: dict, cfg: dict) -> list[str]:
    records = kb.get("records", [])[:18]
    meta = kb.get("meta", {})
    doc_title = meta.get("doc_title") or meta.get("source") or ""
    fallback = _fallback_example_queries(records, doc_title)
    if not records:
        return fallback

    items = "\n".join(
        f"- [{r.get('section_id')}] {r.get('title')}: {(r.get('summary') or r.get('details') or '')[:260]}"
        for r in records
    )[:5000]
    prompt = f"""สร้างคำถามแนะนำ 6-8 ข้อสำหรับผู้ใช้ถามเอกสารนี้ในระบบ RAG

เงื่อนไข:
- ต้องเฉพาะเจาะจงกับเนื้อหาเอกสาร ไม่ใช้คำถาม generic
- ภาษาไทย กระชับ เป็นคำถามที่ตอบจากเอกสารได้
- ถ้ามีชื่อหน่วยงาน กฎหมาย มาตรา โครงการ วันที่ หรือเงื่อนไข ให้สะท้อนในคำถาม
- ตอบเป็น JSON array ของ string เท่านั้น

ชื่อเอกสาร: {doc_title}

หัวข้อ/สรุป:
{items}
"""
    try:
        raw = await azure_once([
            {"role": "system", "content": "คุณช่วยสร้างคำถามแนะนำจากเอกสาร ตอบ JSON array เท่านั้น"},
            {"role": "user", "content": prompt},
        ], cfg)
        raw = raw.strip()
        if raw.startswith("```"):
            raw = re.sub(r"^```[a-z]*\n?", "", raw)
            raw = re.sub(r"\n?```$", "", raw)
        data = json.loads(raw)
        if isinstance(data, dict):
            data = data.get("questions") or data.get("example_queries") or []
        queries = [str(q).strip() for q in data if str(q).strip()]
        if len(queries) >= 3:
            return queries[:8]
    except Exception as exc:
        print(f"[example_queries] fallback due to: {exc}")
    return fallback


def build_hierarchy(records: list) -> dict:
    chapters: dict = {}
    for rec in records:
        ch = rec["chapter_no"]
        if ch not in chapters:
            chapters[ch] = {"id": ch, "title": f"บทที่ {ch}", "sections": {}}
        sid = rec["section_id"]
        chapters[ch]["sections"][sid] = {"id": sid, "title": rec["title"],
                                          "record_idx": records.index(rec), "children": []}
    for ch_data in chapters.values():
        for sid in list(ch_data["sections"]):
            parts = sid.split(".")
            if len(parts) > 1:
                parent = ".".join(parts[:-1])
                if parent in ch_data["sections"]:
                    ch_data["sections"][parent]["children"].append(sid)
    return chapters

# ─────────────────────────────────────────────────────────────
#  API ROUTES
# ─────────────────────────────────────────────────────────────

@app.get("/api/config")
async def api_get_config():
    cfg = load_cfg()
    return {k: ("***" if k == "api_key" and cfg[k] else cfg[k]) for k in cfg}

@app.post("/api/config")
async def api_save_config(body: dict):
    save_cfg(body)
    _bm25_cache.clear()
    return {"ok": True}

@app.get("/api/docs")
async def api_docs():
    return all_docs()

@app.delete("/api/docs/{doc_id}")
async def api_delete_doc(doc_id: str):
    root = KB_DIR.resolve()
    doc_dir = (KB_DIR / doc_id).resolve()
    if doc_dir.parent != root or not doc_dir.is_dir():
        raise HTTPException(404, "document not found")

    task = _graph_tasks.pop(doc_id, None)
    if task and not task.done():
        task.cancel()
    _bm25_cache.pop(doc_id, None)

    shutil.rmtree(doc_dir)

    async with _sess_lock:
        sess = _load_sessions()
        suffix = f":{doc_id}"
        for key in list(sess.keys()):
            if key.endswith(suffix):
                del sess[key]
        _persist_sessions(sess)

    return {"ok": True, "doc_id": doc_id}


_REEXTRACT_ARTIFACTS = {
    "knowledge_base.json",
    "document_meta.json",
    "graph.json",
    "source_clean.json",
    "source.txt",
}

_REEXTRACT_META_KEYS = {
    "document_title",
    "document_type",
    "document_no",
    "published_date",
    "effective_date",
    "agency",
    "ministry",
    "related_agencies",
    "legal_domain",
    "collection",
    "source_type",
    "project_type",
    "relationship_type",
    "risk_flags",
    "last_reviewed",
    "procurement_method",
    "procurement_status",
    "work_category",
    "page_count",
    "file_size_bytes",
    "source_domain",
    "metadata_schema_version",
    "extracted_by",
    "graph_status",
}


async def reset_extraction_artifacts(doc_id: str) -> dict:
    doc_dir = KB_DIR / doc_id
    if not doc_dir.exists() or not doc_dir.is_dir():
        raise HTTPException(404, "document not found")

    task = _graph_tasks.pop(doc_id, None)
    if task and not task.done():
        task.cancel()
    _bm25_cache.pop(doc_id, None)

    removed = []
    for name in _REEXTRACT_ARTIFACTS:
        path = doc_dir / name
        if path.exists() and path.is_file():
            path.unlink()
            removed.append(name)

    meta = load_meta(doc_id) or {"doc_id": doc_id}
    for key in _REEXTRACT_META_KEYS:
        meta.pop(key, None)
    meta["status"] = "uploaded"
    meta["total_records"] = 0
    meta["reextract_requested_at"] = time.time()
    save_meta(doc_id, meta)

    async with _sess_lock:
        sess = _load_sessions()
        suffix = f":{doc_id}"
        for key in list(sess.keys()):
            if key.endswith(suffix):
                del sess[key]
        _persist_sessions(sess)

    return {"doc_id": doc_id, "removed": removed}


@app.get("/api/docs/{doc_id}/index")
async def api_doc_index(doc_id: str):
    kb = load_kb(doc_id)
    if not kb: raise HTTPException(404)
    examples = kb.get("example_queries") or _fallback_example_queries(
        kb.get("records", []),
        (kb.get("meta") or {}).get("doc_title") or (kb.get("meta") or {}).get("source") or "",
    )
    return {
        "meta": kb.get("meta"),
        "page_index": kb.get("page_index"),
        "example_queries": examples,
    }

@app.get("/api/docs/{doc_id}/metadata")
async def api_doc_metadata(doc_id: str):
    meta = load_meta(doc_id)
    if not meta:
        raise HTTPException(404, "document not found")
    return {
        "doc_id": doc_id,
        "meta": meta,
        "document_meta": load_document_meta(doc_id),
    }

@app.post("/api/search")
async def api_search(req: SearchReq):
    q = (req.query or "").strip()
    if not q:
        raise HTTPException(400, "query required")
    k = max(1, min(req.k, 20))
    allowed_doc_ids = {
        doc["doc_id"]
        for doc in all_docs()
        if doc.get("status") == "ready"
        and _document_visible(doc, req.user_id, req.search_scope)
    }
    if not req.doc_id:
        return graph_first_search(q, k=k, allowed_doc_ids=allowed_doc_ids)

    if req.doc_id not in allowed_doc_ids:
        raise HTTPException(404, "document not found")

    candidates = [req.doc_id] if req.doc_id else [d["doc_id"] for d in all_docs() if d.get("status") == "ready"]
    hits = []
    for doc_id in candidates:
        bm25 = get_bm25(doc_id)
        if not bm25:
            continue
        meta = load_meta(doc_id) or {}
        for rec, score in bm25.search(q, k=k):
            hits.append({
                "doc_id": doc_id,
                "doc_title": meta.get("filename") or doc_id,
                "source_kind": meta.get("source_kind"),
                "source_origin": meta.get("source_origin"),
                "source_url": meta.get("source_url"),
                "section_id": rec.get("section_id"),
                "title": rec.get("title"),
                "summary": rec.get("summary"),
                "details": rec.get("details", "")[:1200],
                "keywords": rec.get("keywords") or [],
                "score": score,
            })
    hits.sort(key=lambda item: item["score"], reverse=True)
    return {"query": q, "hits": hits[:k]}


@app.get("/api/graph")
async def api_global_graph():
    return {"ready": True, "status": "ready", "graph": prune_graph_for_ui(build_global_graph())}


@app.post("/api/graph/search")
async def api_graph_search(req: GraphSearchReq):
    q = (req.query or "").strip()
    if not q:
        raise HTTPException(400, "query required")
    return graph_first_search(q, k=max(1, min(req.k, 30)), doc_limit=max(1, min(req.doc_limit, 20)))

@app.get("/api/docs/{doc_id}/sections/{section_id:path}")
async def api_doc_section(doc_id: str, section_id: str):
    kb = load_kb(doc_id)
    if not kb:
        raise HTTPException(404, "Knowledge base not ready")
    for rec in kb.get("records", []):
        if str(rec.get("section_id")) == section_id:
            return {"doc_id": doc_id, "record": rec, "meta": kb.get("meta") or {}}
    raise HTTPException(404, "section not found")

@app.get("/api/docs/{doc_id}/source")
async def api_doc_source(doc_id: str):
    meta = load_meta(doc_id)
    if not meta:
        raise HTTPException(404, "document not found")
    doc_dir = KB_DIR / doc_id
    source_text = ""
    for name in ("source.txt", "source.url"):
        p = doc_dir / name
        if p.exists():
            source_text = p.read_text(encoding="utf-8")[:20000]
            break
    return {"doc_id": doc_id, "meta": meta, "source_preview": source_text}

@app.get("/api/docs/{doc_id}/graph")
async def api_doc_graph(doc_id: str):
    graph = load_graph(doc_id)
    running = _graph_running(doc_id)
    if graph and not running:
        return {"ready": True, "status": "ready", "graph": graph}

    meta = load_meta(doc_id) or {}
    status = "processing" if running else meta.get("graph_status") or "missing"
    if status == "ready" and graph:
        return {"ready": True, "status": "ready", "graph": graph}
    if status == "error":
        return {
            "ready": False,
            "status": "error",
            "doc_id": doc_id,
            "msg": meta.get("graph_error") or "Graph extraction failed",
        }
    return {
        "ready": False,
        "status": status,
        "doc_id": doc_id,
        "msg": "กำลังวิเคราะห์กราฟความสัมพันธ์...",
    }

@app.post("/api/docs/{doc_id}/graph")
async def api_build_doc_graph(doc_id: str, force: bool = False):
    kb = load_kb(doc_id)
    if not kb:
        raise HTTPException(404, "Knowledge base not ready")
    cfg = load_cfg()
    if not cfg.get("api_key"):
        raise HTTPException(400, "API key not configured")

    graph = load_graph(doc_id)
    if graph and not force and not _graph_running(doc_id):
        return {"ready": True, "status": "ready", "graph": graph}

    started = schedule_graph_build(doc_id, cfg, force=force)
    return {
        "ready": False,
        "status": "processing",
        "started": started,
        "doc_id": doc_id,
        "msg": "กำลังวิเคราะห์กราฟความสัมพันธ์ด้วย Hyper-Extract...",
    }

@app.post("/api/upload")
async def api_upload(file: UploadFile = File(...)):
    doc_id  = uuid.uuid4().hex[:8]
    doc_dir = KB_DIR / doc_id
    doc_dir.mkdir()
    pdf_path = doc_dir / "source.pdf"
    pdf_path.write_bytes(await file.read())
    meta = {"doc_id": doc_id, "filename": file.filename, "source_kind": "pdf", "source_origin": "manual_upload",
            "status": "uploaded", "created": time.time(), "total_records": 0}
    save_meta(doc_id, meta)
    source = build_source_metadata(
        doc_id=doc_id,
        source_kind="pdf",
        source_origin="manual_upload",
        source_path=pdf_path,
        original_filename=file.filename,
    )
    save_initial_document_metadata(doc_id, source)
    return {"doc_id": doc_id, "filename": file.filename}


def _document_visible(meta: dict, user_id: str | None, search_scope: str) -> bool:
    owner = str(meta.get("owner_id") or "").strip() or None
    viewer = str(user_id or "").strip() or None
    scope = search_scope if search_scope in {"global", "mine", "both"} else "both"
    if scope == "global":
        return owner is None
    if scope == "mine":
        return owner is not None and viewer is not None and owner == viewer
    return owner is None or (viewer is not None and owner == viewer)


def _document_text(file_bytes: bytes, suffix: str) -> str:
    if suffix == ".txt":
        for encoding in ("utf-8-sig", "utf-8", "tis-620"):
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("ไม่สามารถอ่าน encoding ของไฟล์ TXT ได้")
    if suffix == ".docx":
        from docx import Document

        document = Document(io.BytesIO(file_bytes))
        blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                line = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if line:
                    blocks.append(line)
        return "\n\n".join(blocks)
    raise ValueError(f"unsupported document type: {suffix}")


@app.post("/api/ingest")
async def api_ingest(
    file: UploadFile = File(...),
    doc_id: Optional[str] = Form(default=None),
    display_name: Optional[str] = Form(default=None),
    category: Optional[str] = Form(default=None),
    owner_id: Optional[str] = Form(default=None),
    scope: str = Form(default="baseline"),
    replace: bool = Form(default=False),
):
    """Ingest a TOR document into PageIndex and return only after it is searchable."""
    requested_id = (doc_id or uuid.uuid4().hex[:8]).strip()
    if not re.fullmatch(r"[A-Za-z0-9_-]{1,64}", requested_id):
        raise HTTPException(400, "invalid doc_id")

    original_name = file.filename or "document.pdf"
    suffix = Path(original_name).suffix.lower()
    if suffix not in {".pdf", ".docx", ".txt"}:
        raise HTTPException(400, "PageIndex supports PDF, DOCX, and TXT")
    content = await file.read()
    if not content:
        raise HTTPException(400, "empty document")

    doc_dir = KB_DIR / requested_id
    if doc_dir.exists() and not replace:
        raise HTTPException(409, "document already exists")

    stage_dir = KB_DIR / f".ingest-{requested_id}-{uuid.uuid4().hex[:8]}"
    stage_dir.mkdir()
    source_path = stage_dir / f"source{suffix}"
    source_path.write_bytes(content)
    cfg = load_cfg()
    if not cfg.get("api_key"):
        shutil.rmtree(stage_dir, ignore_errors=True)
        raise HTTPException(400, "PageIndex LLM is not configured")

    try:
        from pageindex_extractor import extract_text_with_pageindex, extract_with_pageindex

        if suffix == ".pdf":
            kb = await extract_with_pageindex(str(source_path), requested_id, cfg)
        else:
            text = _document_text(content, suffix)
            if not text.strip():
                raise ValueError("เอกสารไม่มีข้อความสำหรับสร้าง PageIndex")
            kb = await extract_text_with_pageindex(
                text,
                requested_id,
                cfg,
                source_name=display_name or original_name,
                source_path=str(source_path),
                source_kind=suffix.lstrip("."),
            )

        records = kb.get("records") or []
        source = build_source_metadata(
            doc_id=requested_id,
            source_kind=suffix.lstrip("."),
            source_origin="tor_generator",
            source_path=source_path,
            original_filename=original_name,
        )
        save_document_metadata(stage_dir, empty_document_metadata(source))
        (stage_dir / "knowledge_base.json").write_text(
            json.dumps(kb, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        meta = {
            "doc_id": requested_id,
            "filename": display_name or original_name,
            "original_filename": original_name,
            "source_kind": suffix.lstrip("."),
            "source_origin": "tor_generator",
            "status": "ready",
            "created": time.time(),
            "total_records": len(records),
            "extracted_by": "pageindex",
            "category": category,
            "owner_id": owner_id or None,
            "scope": scope,
        }
        (stage_dir / "meta.json").write_text(
            json.dumps(meta, ensure_ascii=False), encoding="utf-8"
        )

        if doc_dir.exists():
            shutil.rmtree(doc_dir)
        stage_dir.replace(doc_dir)
        _bm25_cache.pop(requested_id, None)
        return {
            "ok": True,
            "doc_id": requested_id,
            "status": "ready",
            "total_records": len(records),
            "engine": "pageindex",
        }
    except HTTPException:
        shutil.rmtree(stage_dir, ignore_errors=True)
        raise
    except Exception as exc:
        shutil.rmtree(stage_dir, ignore_errors=True)
        raise HTTPException(500, f"PageIndex ingestion failed: {exc}") from exc

@app.post("/api/upload-link")
async def api_upload_link(req: LinkUploadReq):
    from urllib.parse import urlparse

    url = (req.url or "").strip()
    if not url:
        raise HTTPException(400, "URL required")
    if re.search(r"\s", url):
        raise HTTPException(400, "Invalid URL")
    if not re.match(r"^https?://", url, re.I):
        url = "https://" + url
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https") or not parsed.netloc:
        raise HTTPException(400, "Invalid URL")

    doc_id = uuid.uuid4().hex[:8]
    doc_dir = KB_DIR / doc_id
    doc_dir.mkdir()
    (doc_dir / "source.url").write_text(url, encoding="utf-8")
    label = parsed.netloc + (parsed.path if parsed.path and parsed.path != "/" else "")
    meta = {
        "doc_id": doc_id,
        "filename": label[:160],
        "source_kind": "link",
        "source_origin": "manual_upload",
        "source_url": url,
        "status": "uploaded",
        "created": time.time(),
        "total_records": 0,
    }
    save_meta(doc_id, meta)
    source = build_source_metadata(
        doc_id=doc_id,
        source_kind="link",
        source_origin="manual_upload",
        source_url=url,
        original_filename=meta["filename"],
    )
    save_initial_document_metadata(doc_id, source)
    return {"doc_id": doc_id, "filename": meta["filename"], "source_url": url}

@app.get("/api/line/messages")
async def api_line_messages(status: str = "pending", limit: int = 50, offset: int = 0):
    params = {"limit": max(1, min(limit, 100)), "offset": max(0, offset)}
    if status:
        params["status"] = status
    try:
        async with httpx.AsyncClient(timeout=20) as client:
            resp = await client.get(f"{LINE_HUB_API_BASE}/messages", params=params)
            resp.raise_for_status()
            return resp.json()
    except Exception as exc:
        raise HTTPException(502, f"LINE hub unavailable: {exc}") from exc

@app.post("/api/line/import")
async def api_line_import(req: LineImportReq):
    try:
        item = req.message or {}
        if not item and req.message_id:
            item = await fetch_line_message(req.message_id)
        if not item:
            raise ValueError("message or message_id required")
        return create_line_link_doc(item, source="pull")
    except ValueError as exc:
        raise HTTPException(400, str(exc)) from exc
    except Exception as exc:
        raise HTTPException(502, f"Cannot import LINE message: {exc}") from exc

@app.post("/api/line/ingest")
async def api_line_ingest(payload: dict):
    item = {
        "id": payload.get("external_id") or payload.get("id"),
        "sender": payload.get("sender"),
        "source": payload.get("source"),
        "receivedAt": payload.get("received_at") or payload.get("receivedAt"),
        "status": "pending",
        "message": {
            "id": payload.get("message_id"),
            "type": payload.get("content_type") or "text",
            "content": payload.get("content"),
            "format": payload.get("content_format"),
            "links": payload.get("links") or [],
            "url": payload.get("url"),
            "markdown": payload.get("markdown"),
        },
        "metadata": payload.get("metadata") or {},
    }
    try:
        result = create_line_link_doc(item, source="push")
        return {"ok": True, **result}
    except ValueError as exc:
        raise HTTPException(400, str(exc)) from exc

@app.get("/api/extract/{doc_id}")
async def api_extract(doc_id: str, engine: str = "default", reset: bool = False):
    """Extract a PDF into a knowledge base.

    Query param:
        engine=default    — original section-based pipeline (pdfplumber + Azure enrich)
        engine=pageindex  — VectifyAI/PageIndex hierarchical tree pipeline
    """
    doc_dir = KB_DIR / doc_id
    if not doc_dir.exists(): raise HTTPException(404)
    cfg = load_cfg()
    if not cfg.get("api_key"): raise HTTPException(400, "API key not configured")

    async def stream():
        def evt(event: str, data: dict) -> str:
            return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"

        stream_started = time.time()
        last_progress = {"pct": 0, "msg": "กำลังเริ่ม extract", "updated": stream_started}

        def remember_progress(payload: dict):
            if not isinstance(payload, dict):
                return
            last_progress["pct"] = max(float(last_progress.get("pct") or 0), float(payload.get("pct") or 0))
            if payload.get("msg"):
                last_progress["msg"] = str(payload["msg"])
            last_progress["updated"] = time.time()

        def heartbeat_payload(phase: str) -> dict:
            idle = int(time.time() - float(last_progress.get("updated") or stream_started))
            total = int(time.time() - stream_started)
            return {
                "pct": last_progress.get("pct") or 1,
                "msg": f"{phase}: {last_progress.get('msg')} · ยังทำงานอยู่ {idle}s · รวม {total}s",
            }

        meta0 = load_meta(doc_id) or {}
        if not meta0.get("source_kind"):
            if (doc_dir / "source.url").exists():
                meta0["source_kind"] = "link"
            elif (doc_dir / "source.pdf").exists():
                meta0["source_kind"] = "pdf"

        if meta0.get("source_kind") == "link":
            source_url_ok = bool(meta0.get("source_url")) or (doc_dir / "source.url").exists()
            if not source_url_ok:
                yield evt("error", {"msg": "Re-extract ไม่ได้: ไม่พบ source.url หรือ source_url เดิมของเอกสารนี้"})
                return
        elif not (doc_dir / "source.pdf").exists():
            yield evt("error", {"msg": "Re-extract ไม่ได้: ไม่พบ source.pdf เดิมของเอกสารนี้"})
            return

        if reset:
            payload = {"pct": 1, "msg": "กำลังลบข้อมูล extract เก่า..."}
            remember_progress(payload)
            yield evt("progress", payload)
            reset_info = await reset_extraction_artifacts(doc_id)
            removed = len(reset_info.get("removed") or [])
            payload = {"pct": 3, "msg": f"ลบข้อมูลเก่าแล้ว {removed} ไฟล์ · กำลัง extract ใหม่..."}
            remember_progress(payload)
            yield evt("progress", payload)
            meta0 = load_meta(doc_id) or meta0

        if meta0.get("source_kind") == "link":
            try:
                from link_extractor import ai_clean_article_text, fetch_url_text
                from pageindex_extractor import extract_text_with_pageindex
            except ImportError as e:
                yield evt("error", {"msg": f"link extractor ไม่พร้อม: {e}"})
                return

            q: asyncio.Queue = asyncio.Queue()

            async def progress_cb(pct: int, msg: str):
                await q.put(("progress", {"pct": pct, "msg": msg}))

            async def run_link_extraction():
                try:
                    source_url = meta0.get("source_url") or (doc_dir / "source.url").read_text(encoding="utf-8").strip()
                    await progress_cb(5, "กำลัง crawl ลิงก์ข่าว...")
                    raw = await fetch_url_text(source_url)
                    await progress_cb(20, "ดึงหน้าเว็บสำเร็จ · กำลังให้ AI กรองเนื้อข่าวจริง...")
                    cleaned = await ai_clean_article_text(raw, cfg)
                    content = cleaned.get("content") or raw.get("text") or ""
                    if len(content.strip()) < 200:
                        raise ValueError("ไม่พบเนื้อหาข่าวที่เพียงพอจากลิงก์นี้")

                    title = cleaned.get("title") or raw.get("title") or meta0.get("filename") or source_url
                    source_txt = doc_dir / "source.txt"
                    source_txt.write_text(content, encoding="utf-8")
                    (doc_dir / "source_clean.json").write_text(
                        json.dumps(
                            {
                                "url": raw.get("url") or source_url,
                                "title": title,
                                "published_date": cleaned.get("published_date"),
                                "source_name": cleaned.get("source_name"),
                            },
                            ensure_ascii=False,
                            indent=2,
                        ),
                        encoding="utf-8",
                    )
                    await progress_cb(35, "เตรียม plain text สำหรับ PageIndex...")
                    kb = await extract_text_with_pageindex(
                        content,
                        doc_id,
                        cfg,
                        progress_cb,
                        source_name=title,
                        source_url=raw.get("url") or source_url,
                        source_path=str(source_txt),
                    )
                    await q.put(("_kb", {"kb": kb, "title": title, "source_url": raw.get("url") or source_url}))
                except Exception as exc:
                    await q.put(("error", {"msg": str(exc)}))
                finally:
                    await q.put(("_done", None))

            task = asyncio.create_task(run_link_extraction())
            kb_payload = None
            while True:
                try:
                    event_name, payload = await asyncio.wait_for(q.get(), timeout=8)
                except asyncio.TimeoutError:
                    payload = heartbeat_payload("Link/PageIndex")
                    remember_progress(payload)
                    yield evt("progress", payload)
                    continue
                if event_name == "_done":
                    break
                if event_name == "_kb":
                    kb_payload = payload
                    continue
                if event_name == "progress":
                    remember_progress(payload)
                yield evt(event_name, payload)
                if event_name == "error":
                    task.cancel()
                    return

            if not kb_payload:
                yield evt("error", {"msg": "ไม่สามารถสร้าง Knowledge Base จากลิงก์นี้ได้"})
                return

            kb = kb_payload["kb"]
            payload = {"pct": 92, "msg": "AI กำลังสกัด metadata กลางจากเนื้อหา..."}
            remember_progress(payload)
            yield evt("progress", payload)
            source = build_source_metadata(
                doc_id=doc_id,
                source_kind="link",
                source_origin=(load_meta(doc_id) or {}).get("source_origin") or "manual_upload",
                source_path=doc_dir / "source.txt",
                original_filename=kb_payload["title"],
                source_url=kb_payload["source_url"],
            )
            metadata_task = asyncio.create_task(enrich_and_save_document_metadata(
                doc_id,
                (doc_dir / "source.txt").read_text(encoding="utf-8"),
                cfg,
                source=source,
                fallback_title=kb_payload["title"],
                kb=kb,
            ))
            metadata_started = time.time()
            while not metadata_task.done():
                try:
                    await asyncio.wait_for(asyncio.shield(metadata_task), timeout=8)
                except asyncio.TimeoutError:
                    elapsed = int(time.time() - metadata_started)
                    yield evt("progress", {"pct": 92, "msg": f"AI กำลังสกัด metadata กลางจากเนื้อหา... · {elapsed}s"})
            await metadata_task

            payload = {"pct": 95, "msg": "AI กำลังสร้างคำถามแนะนำจากเนื้อหานี้..."}
            remember_progress(payload)
            yield evt("progress", payload)
            query_task = asyncio.create_task(generate_example_queries(kb, cfg))
            query_started = time.time()
            while not query_task.done():
                try:
                    kb["example_queries"] = await asyncio.wait_for(asyncio.shield(query_task), timeout=8)
                except asyncio.TimeoutError:
                    elapsed = int(time.time() - query_started)
                    yield evt("progress", {"pct": 95, "msg": f"AI กำลังสร้างคำถามแนะนำจากเนื้อหานี้... · {elapsed}s"})
            if "example_queries" not in kb:
                kb["example_queries"] = await query_task

            payload = {"pct": 97, "msg": "กำลังบันทึก Knowledge Base จากลิงก์..."}
            remember_progress(payload)
            yield evt("progress", payload)
            (doc_dir / "knowledge_base.json").write_text(
                json.dumps(kb, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            records = kb.get("records", [])
            meta = load_meta(doc_id) or {}
            meta.update({
                "filename": meta.get("filename") or kb_payload["title"],
                "source_url": kb_payload["source_url"],
                "source_kind": "link",
                "source_origin": meta.get("source_origin") or "manual_upload",
                "status": "ready",
                "total_records": len(records),
                "extracted_by": "pageindex",
            })
            save_meta(doc_id, meta)
            if meta.get("source_origin") == "line_api" and meta.get("external_id"):
                try:
                    await ack_line_message(meta["external_id"])
                    meta["line_status"] = "acknowledged"
                    meta["line_acknowledged_at"] = time.time()
                    meta.pop("line_ack_error", None)
                    save_meta(doc_id, meta)
                except Exception as exc:
                    meta["line_ack_error"] = str(exc)
                    save_meta(doc_id, meta)
                    print(f"[line] ack failed for {meta.get('external_id')}: {exc}")
            _bm25_cache.pop(doc_id, None)
            try:
                schedule_graph_build(doc_id, cfg, force=True)
            except Exception as exc:
                print(f"[graph] schedule failed for {doc_id}: {exc}")

            yield evt("done", {"pct": 100, "doc_id": doc_id, "total": len(records),
                               "msg": f"✅ Knowledge Base จากลิงก์พร้อมแล้ว ({len(records)} sections)"})
            return

        # ── PageIndex engine ─────────────────────────────────────────
        if engine == "pageindex":
            try:
                from pageindex_extractor import extract_with_pageindex
            except ImportError as e:
                yield evt("error", {"msg": f"pageindex_extractor ไม่พร้อม: {e}"})
                return

            async def progress_cb(pct: int, msg: str):
                pass  # progress is collected below via queue

            # Use a queue so we can yield SSE events from inside the coroutine
            q: asyncio.Queue = asyncio.Queue()

            async def progress_cb(pct: int, msg: str):  # noqa: F811
                await q.put(("progress", {"pct": pct, "msg": msg}))

            async def run_extraction():
                try:
                    kb = await extract_with_pageindex(
                        str(doc_dir / "source.pdf"),
                        doc_id, cfg, progress_cb,
                    )
                    await q.put(("_kb", kb))
                except Exception as exc:
                    await q.put(("error", {"msg": str(exc)}))
                finally:
                    await q.put(("_done", None))

            task = asyncio.create_task(run_extraction())
            kb = None
            while True:
                try:
                    event_name, payload = await asyncio.wait_for(q.get(), timeout=8)
                except asyncio.TimeoutError:
                    payload = heartbeat_payload("PageIndex")
                    remember_progress(payload)
                    yield evt("progress", payload)
                    continue
                if event_name == "_done":
                    break
                if event_name == "_kb":
                    kb = payload
                    continue
                if event_name == "progress":
                    remember_progress(payload)
                yield evt(event_name, payload)
                if event_name == "error":
                    task.cancel()
                    return

            if kb is None:
                yield evt("error", {"msg": "PageIndex ไม่ได้คืนค่า knowledge_base"})
                return

            # Save
            payload = {"pct": 93, "msg": "AI กำลังสกัด metadata กลางจากเอกสาร..."}
            remember_progress(payload)
            yield evt("progress", payload)
            record_text = "\n\n".join(
                " ".join(filter(None, [
                    r.get("title", ""),
                    r.get("summary", ""),
                    r.get("details", ""),
                    " ".join(r.get("keywords") or []),
                ]))
                for r in kb.get("records", [])[:24]
            )
            source = build_source_metadata(
                doc_id=doc_id,
                source_kind="pdf",
                source_origin=(load_meta(doc_id) or {}).get("source_origin") or "manual_upload",
                source_path=doc_dir / "source.pdf",
                original_filename=(load_meta(doc_id) or {}).get("filename"),
            )
            metadata_task = asyncio.create_task(enrich_and_save_document_metadata(
                doc_id,
                record_text,
                cfg,
                source=source,
                fallback_title=(kb.get("meta") or {}).get("doc_title") or (load_meta(doc_id) or {}).get("filename"),
                kb=kb,
            ))
            metadata_started = time.time()
            while not metadata_task.done():
                try:
                    await asyncio.wait_for(asyncio.shield(metadata_task), timeout=8)
                except asyncio.TimeoutError:
                    elapsed = int(time.time() - metadata_started)
                    yield evt("progress", {"pct": 93, "msg": f"AI กำลังสกัด metadata กลางจากเอกสาร... · {elapsed}s"})
            await metadata_task

            payload = {"pct": 95, "msg": "AI กำลังสร้างคำถามแนะนำจากเอกสารนี้..."}
            remember_progress(payload)
            yield evt("progress", payload)
            query_task = asyncio.create_task(generate_example_queries(kb, cfg))
            query_started = time.time()
            while not query_task.done():
                try:
                    kb["example_queries"] = await asyncio.wait_for(asyncio.shield(query_task), timeout=8)
                except asyncio.TimeoutError:
                    elapsed = int(time.time() - query_started)
                    yield evt("progress", {"pct": 95, "msg": f"AI กำลังสร้างคำถามแนะนำจากเอกสารนี้... · {elapsed}s"})
            if "example_queries" not in kb:
                kb["example_queries"] = await query_task

            payload = {"pct": 97, "msg": "กำลังบันทึก Knowledge Base..."}
            remember_progress(payload)
            yield evt("progress", payload)
            (doc_dir / "knowledge_base.json").write_text(
                json.dumps(kb, ensure_ascii=False, indent=2)
            )
            records = kb.get("records", [])
            meta = load_meta(doc_id) or {}
            meta.update({"status": "ready", "total_records": len(records),
                         "extracted_by": "pageindex"})
            save_meta(doc_id, meta)
            _bm25_cache.pop(doc_id, None)
            try:
                schedule_graph_build(doc_id, cfg, force=True)
            except Exception as exc:
                print(f"[graph] schedule failed for {doc_id}: {exc}")

            yield evt("done", {"pct": 100, "doc_id": doc_id, "total": len(records),
                               "msg": f"✅ Knowledge Base พร้อมแล้ว ({len(records)} sections) [PageIndex]"})
            return

        # ── Default engine (original pipeline) ──────────────────────
        yield evt("progress", {"pct": 5,  "msg": "กำลังอ่าน PDF (Azure Document Intelligence)..."})
        pages = await extract_pages(str(doc_dir / "source.pdf"), cfg)
        yield evt("progress", {"pct": 12, "msg": f"อ่านได้ {len(pages)} หน้า · กำลังสกัด metadata กลาง..."})

        # ── Step 1: Classify document (doc-level metadata) ──
        source = build_source_metadata(
            doc_id=doc_id,
            source_kind="pdf",
            source_origin=(load_meta(doc_id) or {}).get("source_origin") or "manual_upload",
            source_path=doc_dir / "source.pdf",
            original_filename=(load_meta(doc_id) or {}).get("filename"),
        )
        metadata = await enrich_and_save_document_metadata(
            doc_id,
            "\n\n".join(p["text"] for p in pages[:8]),
            cfg,
            source=source,
            fallback_title=(load_meta(doc_id) or {}).get("filename"),
        )
        doc_meta = flatten_for_kb(metadata)
        yield evt("progress", {"pct": 20, "msg": f"จำแนกได้: {doc_meta.get('collection','?')} / {doc_meta.get('source_type','?')} · กำลัง detect sections..."})

        sections = parse_sections(pages)
        if not sections:
            yield evt("error", {"msg": "ไม่พบ sections — ลอง PDF อื่น"})
            return
        yield evt("progress", {"pct": 28, "msg": f"พบ {len(sections)} sections · กำลัง AI enrich..."})

        # ── Step 2: Enrich sections in batches ──
        BATCH = 5
        enriched: list[dict] = []
        for start in range(0, len(sections), BATCH):
            batch = sections[start:start + BATCH]
            data  = await enrich_batch(batch, cfg)
            enriched.extend(data)
            pct = 28 + int(65 * (start + BATCH) / len(sections))
            yield evt("progress", {"pct": min(pct, 92),
                                   "msg": f"Enriched {min(start+BATCH, len(sections))}/{len(sections)} sections..."})

        # ── Step 3: Build records ──
        records = []
        for i, (sec, ai) in enumerate(zip(sections, enriched)):
            ch = sec["section_id"].split(".")[0]
            sec_ld    = ai.get("legal_domain") or doc_meta.get("legal_domain")
            sec_topic = list(set((ai.get("topic") or []) + (doc_meta.get("topic") or [])))
            sec_uc    = list(set((ai.get("use_case") or []) + (doc_meta.get("use_case") or [])))

            rec = {
                "id":         f"DOC-{i+1:04d}",
                "section_id": sec["section_id"],
                "chapter_no": ch,
                "chapter":    f"บทที่ {ch}",
                "title":      sec["title"],
                "summary":    ai.get("summary") or sec["title"],
                "details":    ai.get("details") or sec["content"][:600],
                "keywords":   ai.get("keywords") or [],
                "section_meta": {
                    "legal_domain": sec_ld,
                    "topic":        sec_topic,
                    "use_case":     sec_uc,
                    "case_id":      None,
                },
                "page_start": sec["page"],
                "page_end":   sec["page"],
                "boost":      float(ai.get("boost", 1.0)),
                "search_text": " ".join(filter(None, [
                    sec["section_id"], sec["title"],
                    ai.get("summary", ""), sec["content"],
                    " ".join(ai.get("keywords", [])),
                    doc_meta.get("ministry") or "",
                    doc_meta.get("agency") or "",
                    " ".join(doc_meta.get("topic") or []),
                    doc_meta.get("legal_domain") or "",
                ]))
            }
            records.append(rec)

        meta = load_meta(doc_id) or {}
        if doc_meta.get("document_title"):
            doc_meta_copy = {k: v for k, v in doc_meta.items() if k != "document_title"}
        else:
            doc_meta_copy = doc_meta

        kb = {
            "meta": {
                "doc_title":    doc_meta.get("document_title") or meta.get("filename", "Doc"),
                "source":       meta.get("filename", ""),
                "total_records": len(records),
                "doc_meta":     doc_meta_copy,
            },
            "page_index":    build_hierarchy(records),
            "records":       records,
            "synonyms":      {},
            "example_queries": []
        }
        yield evt("progress", {"pct": 94, "msg": "กำลังสร้างคำถามแนะนำจากเอกสารนี้..."})
        kb["example_queries"] = await generate_example_queries(kb, cfg)
        yield evt("progress", {"pct": 96, "msg": "กำลังบันทึก Knowledge Base..."})
        (doc_dir / "knowledge_base.json").write_text(json.dumps(kb, ensure_ascii=False, indent=2))

        meta.update({"status": "ready", "total_records": len(records)})
        save_meta(doc_id, meta)
        _bm25_cache.pop(doc_id, None)

        yield evt("done", {"pct": 100, "doc_id": doc_id, "total": len(records),
                           "msg": f"✅ Knowledge Base พร้อมแล้ว ({len(records)} sections)"})

    return StreamingResponse(stream(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


# ─── Chat ─────────────────────────────────────────────────────

SYSTEM_PROMPT = """คุณเป็นผู้เชี่ยวชาญด้านการจัดซื้อจัดจ้างภาครัฐของไทย ช่วยตอบคำถามโดยอ้างอิงจากเอกสารที่ให้มา

หลักการ:
- ตอบเป็นภาษาไทย กระชับ ชัดเจน อ่านง่าย
- ใช้เฉพาะข้อมูลในบล็อก "อ้างอิง #..." ที่ให้มาเท่านั้น
- ถ้ากล่าวอ้างข้อมูลเฉพาะ ให้ใส่เลขอ้างอิงท้ายประโยค เช่น [อ้างอิง #2]
- ห้ามอ้างหัวข้อหรือเอกสารที่ไม่ได้อยู่ในบล็อกอ้างอิงที่ให้มา
- อ้างอิงแบบคนอ่านเข้าใจ เช่น "เอกสารแนวทางปฏิบัติสำหรับการจ้างทำของ หัวข้อ 3.2" หรือ "หัวข้อหลักเกณฑ์ราคากลาง"
- ห้ามใส่รหัสภายในของระบบหรือ doc_id ในคำตอบ
- ห้ามอ้างเป็นวงเล็บเทคนิคที่มีรหัสเอกสารกับเลขหัวข้อ ให้แปลงเป็นชื่อเอกสารหรือชื่อหัวข้อแทน
- ใช้ bullet points เมื่อมีหลายประเด็น
- ถ้าข้อมูลไม่เพียงพอ บอกตรงๆ และแนะนำหัวข้อที่ควรอ่านเพิ่มเติม"""

class ChatReq(BaseModel):
    doc_id:     Optional[str] = None
    session_id: str        # client-generated UUID, persisted in localStorage
    question:   str


@app.post("/api/chat")
async def api_chat(req: ChatReq):
    chat_doc_id = (req.doc_id or "global").strip() or "global"
    is_global = chat_doc_id == "global"

    def make_source(item: dict, ref_no: int) -> dict:
        return {
            "ref_no": ref_no,
            "doc_id": item.get("doc_id"),
            "doc_title": item.get("doc_title"),
            "section_id": item.get("section_id"),
            "title": item.get("title"),
            "score": round(float(item.get("score") or 0), 2),
        }

    def make_context_block(item: dict, ref_no: int, include_doc: bool = True) -> str:
        doc_line = f"เอกสาร: {item.get('doc_title')}\n" if include_doc else ""
        return (
            f"อ้างอิง #{ref_no}\n"
            f"{doc_line}"
            f"หัวข้อ {item.get('section_id')}: {item.get('title')}\n"
            f"{item.get('summary') or ''}\n{(item.get('details') or '')[:900]}"
        )

    if is_global:
        search = graph_first_search(req.question, k=10, doc_limit=8)
        raw_hits = search.get("hits") or []
        seen: dict = {}
        for hit in raw_hits:
            key = f"{hit.get('doc_id')}:{hit.get('section_id')}"
            if key not in seen or hit.get("score", 0) > seen[key]["score"]:
                seen[key] = {
                    "doc_id": hit.get("doc_id"),
                    "doc_title": hit.get("doc_title"),
                    "section_id": hit.get("section_id"),
                    "title": hit.get("title"),
                    "summary": hit.get("summary"),
                    "details": hit.get("details"),
                    "score": round(float(hit.get("score") or 0), 2),
                }
        context_hits = sorted(seen.values(), key=lambda item: item.get("score", 0), reverse=True)[:10]
        shown_hits = context_hits[:6]
        sources = [make_source(hit, i + 1) for i, hit in enumerate(shown_hits)]
        source_meta = {
            "scope": "global",
            "shown_sources": len(sources),
            "total_sources": len(seen),
            "related_docs": len({hit.get("doc_id") for hit in seen.values() if hit.get("doc_id")}),
            "searched_docs": len([d for d in all_docs() if d.get("status") == "ready"]),
        }
        context = "\n\n---\n\n".join(
            make_context_block(hit, i + 1, include_doc=True)
            for i, hit in enumerate(shown_hits)
        )
    else:
        bm25 = get_bm25(chat_doc_id)
        if not bm25:
            raise HTTPException(404, "Knowledge base not ready")

        hits = bm25.search(req.question, k=10)  # over-fetch then deduplicate
        meta = load_meta(chat_doc_id) or {}
        seen: dict = {}
        for r, s in hits:
            sid = r["section_id"]
            if sid not in seen or s > seen[sid]["score"]:
                seen[sid] = {
                    "doc_id": chat_doc_id,
                    "doc_title": _doc_title(meta, chat_doc_id),
                    "section_id": sid,
                    "title": r["title"],
                    "summary": r.get("summary"),
                    "details": r.get("details"),
                    "score": round(s, 2),
                }
        context_hits = sorted(seen.values(), key=lambda item: item.get("score", 0), reverse=True)[:8]
        shown_hits = context_hits[:4]
        sources = [make_source(hit, i + 1) for i, hit in enumerate(shown_hits)]
        source_meta = {
            "scope": "doc",
            "shown_sources": len(sources),
            "total_sources": len(seen),
            "related_docs": 1 if seen else 0,
            "searched_docs": 1,
        }
        context = "\n\n---\n\n".join(
            make_context_block(hit, i + 1, include_doc=False)
            for i, hit in enumerate(shown_hits)
        )

    # Load server-side conversation history
    history = await get_history(req.session_id, chat_doc_id)

    # Save user question to memory
    await push_message(req.session_id, chat_doc_id, "user", req.question)

    # Build message list for LLM
    scope_hint = (
        "โหมดนี้ค้นจากคลังความรู้รวมหลายเอกสาร ให้อ้างด้วยเลข [อ้างอิง #] จากบล็อกข้อมูลที่ให้มาเท่านั้น และใช้ชื่อเอกสาร/ชื่อหัวข้อแบบอ่านง่าย ห้ามอ้างรหัสระบบ"
        if is_global else
        "โหมดนี้ตอบจากเอกสารที่ผู้ใช้เลือกเท่านั้น ให้อ้างด้วยเลข [อ้างอิง #] จากบล็อกข้อมูลที่ให้มาเท่านั้น และใช้ชื่อ/เลขหัวข้อแบบอ่านง่าย ห้ามอ้างรหัสระบบ"
    )
    llm_messages = [{"role": "system", "content": f"{SYSTEM_PROMPT}\n- {scope_hint}"}]
    for m in history[-12:]:            # last 12 turns = ~6 exchanges
        llm_messages.append({"role": m["role"], "content": m["content"]})
    llm_messages.append({"role": "user", "content":
        f"ข้อมูลอ้างอิง:\n{context}\n\nคำถาม: {req.question}"})

    cfg = load_cfg()
    if not cfg.get("api_key"): raise HTTPException(400, "API key not configured")

    async def stream():
        collected: list[str] = []
        yield f"data: {json.dumps({'type':'sources','sources':sources,'source_meta':source_meta}, ensure_ascii=False)}\n\n"
        try:
            async for chunk in azure_stream(llm_messages, cfg):
                d = json.loads(chunk)
                choices = d.get("choices") or []
                if not choices:
                    continue
                token = choices[0].get("delta", {}).get("content") or ""
                if token:
                    collected.append(token)
                    yield f"data: {json.dumps({'type':'token','token':token}, ensure_ascii=False)}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type':'error','msg':str(e)})}\n\n"

        # Persist AI reply to memory
        if collected:
            await push_message(req.session_id, chat_doc_id, "assistant", "".join(collected))

        yield f"data: {json.dumps({'type':'done'})}\n\n"

    return StreamingResponse(stream(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.get("/api/session/{session_id}/{doc_id}")
async def api_get_session(session_id: str, doc_id: str):
    msgs = await get_history(session_id, doc_id)
    return {"messages": msgs}

@app.delete("/api/session/{session_id}/{doc_id}")
async def api_clear_session(session_id: str, doc_id: str):
    await clear_history(session_id, doc_id)
    return {"ok": True}


@app.get("/health")
async def health():
    return {"ok": True, "service": "pageindex-rag", "mode": "backend-only"}


# The merged product has one UI: tor-generator on port 3000. Keep the legacy
# PageIndex static assets dormant unless explicitly enabled for maintenance.
if PAGEINDEX_UI_ENABLED:
    app.mount("/", StaticFiles(directory=str(STATIC_DIR), html=True), name="static")
else:

    @app.get("/", include_in_schema=False)
    async def backend_root():
        return {
            "service": "pageindex-rag",
            "mode": "backend-only",
            "health": "/health",
            "openapi": "/docs",
            "search": "/api/search",
        }
