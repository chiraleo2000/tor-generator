"""Unit tests for the DOCX generator and Thai formatting utilities.

Tests cover:
- Thai date formatting (Gregorian → Buddhist Era)
- Thai/Arabic numeral conversion
- DOCX generation with proper formatting
- Section numbering configuration
- Page setup (margins, font, headers)
"""

import io
from datetime import date, datetime

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.export.docx_generator import (
    BODY_FONT_SIZE,
    FONT_NAME,
    HEADING_FONT_SIZE,
    PAGE_MARGIN,
    TOR_SECTION_LABELS,
    TOR_SECTION_ORDER,
    DOCXGenerator,
    TORContent,
)
from app.export.thai_formatting import (
    buddhist_era_to_gregorian,
    format_currency_thai,
    format_section_number,
    format_thai_date,
    gregorian_to_buddhist_era,
    to_arabic_numerals,
    to_thai_numerals,
)


# =============================================================================
# Thai Formatting Utilities Tests
# =============================================================================


class TestGregorianToBuddhistEra:
    """Tests for Gregorian → Buddhist Era year conversion."""

    def test_standard_conversion(self):
        assert gregorian_to_buddhist_era(2024) == 2567

    def test_year_2000(self):
        assert gregorian_to_buddhist_era(2000) == 2543

    def test_year_1957(self):
        # The year the Buddhist Era offset formula was standardized
        assert gregorian_to_buddhist_era(1957) == 2500

    def test_offset_is_always_543(self):
        for year in [1900, 2000, 2024, 2100]:
            assert gregorian_to_buddhist_era(year) - year == 543


class TestBuddhistEraToGregorian:
    """Tests for Buddhist Era → Gregorian year conversion."""

    def test_standard_conversion(self):
        assert buddhist_era_to_gregorian(2567) == 2024

    def test_round_trip(self):
        for year in [1990, 2000, 2024, 2050]:
            assert buddhist_era_to_gregorian(gregorian_to_buddhist_era(year)) == year


class TestToThaiNumerals:
    """Tests for Arabic → Thai numeral conversion."""

    def test_single_digits(self):
        assert to_thai_numerals("0") == "๐"
        assert to_thai_numerals("1") == "๑"
        assert to_thai_numerals("9") == "๙"

    def test_multi_digit_number(self):
        assert to_thai_numerals("2567") == "๒๕๖๗"

    def test_preserves_non_digit_chars(self):
        assert to_thai_numerals("4.1") == "๔.๑"
        assert to_thai_numerals("hello 123") == "hello ๑๒๓"

    def test_empty_string(self):
        assert to_thai_numerals("") == ""

    def test_no_digits(self):
        assert to_thai_numerals("กขค") == "กขค"


class TestToArabicNumerals:
    """Tests for Thai → Arabic numeral conversion."""

    def test_single_thai_digits(self):
        assert to_arabic_numerals("๐") == "0"
        assert to_arabic_numerals("๑") == "1"
        assert to_arabic_numerals("๙") == "9"

    def test_multi_digit(self):
        assert to_arabic_numerals("๒๕๖๗") == "2567"

    def test_preserves_non_thai_digit_chars(self):
        assert to_arabic_numerals("๔.๑") == "4.1"

    def test_round_trip(self):
        original = "12345"
        assert to_arabic_numerals(to_thai_numerals(original)) == original


class TestFormatThaiDate:
    """Tests for Thai date formatting."""

    def test_basic_date_arabic_numerals(self):
        d = date(2024, 8, 15)
        result = format_thai_date(d, use_thai_numerals=False)
        assert result == "วันที่ 15 สิงหาคม พ.ศ. 2567"

    def test_basic_date_thai_numerals(self):
        d = date(2024, 8, 15)
        result = format_thai_date(d, use_thai_numerals=True)
        assert result == "วันที่ ๑๕ สิงหาคม พ.ศ. ๒๕๖๗"

    def test_january(self):
        d = date(2024, 1, 1)
        result = format_thai_date(d)
        assert "มกราคม" in result
        assert "2567" in result

    def test_december(self):
        d = date(2024, 12, 31)
        result = format_thai_date(d)
        assert "ธันวาคม" in result

    def test_datetime_input(self):
        dt = datetime(2024, 6, 15, 10, 30, 0)
        result = format_thai_date(dt)
        assert "มิถุนายน" in result
        assert "15" in result

    def test_buddhist_era_year_in_output(self):
        d = date(2025, 3, 20)
        result = format_thai_date(d)
        # 2025 + 543 = 2568
        assert "2568" in result


class TestFormatSectionNumber:
    """Tests for section number formatting."""

    def test_integer_arabic(self):
        assert format_section_number(1) == "1"
        assert format_section_number(13) == "13"

    def test_integer_thai(self):
        assert format_section_number(1, use_thai_numerals=True) == "๑"
        assert format_section_number(13, use_thai_numerals=True) == "๑๓"

    def test_hierarchical_arabic(self):
        assert format_section_number("4.1") == "4.1"

    def test_hierarchical_thai(self):
        assert format_section_number("4.1", use_thai_numerals=True) == "๔.๑"


class TestFormatCurrencyThai:
    """Tests for Thai currency formatting."""

    def test_basic_format(self):
        assert format_currency_thai(1_000_000) == "1,000,000 บาท"

    def test_thai_numerals(self):
        result = format_currency_thai(5_000_000, use_thai_numerals=True)
        assert result == "๕,๐๐๐,๐๐๐ บาท"

    def test_zero(self):
        assert format_currency_thai(0) == "0 บาท"

    def test_small_amount(self):
        assert format_currency_thai(500) == "500 บาท"


# =============================================================================
# DOCX Generator Tests
# =============================================================================


class TestDOCXGeneratorBasic:
    """Tests for basic DOCX generation."""

    def setup_method(self):
        self.generator = DOCXGenerator()

    def test_generate_returns_bytes(self):
        content = TORContent(project_name="Test Project")
        result = self.generator.generate(content)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_generated_bytes_is_valid_docx(self):
        content = TORContent(project_name="ทดสอบ")
        result = self.generator.generate(content)
        # Should be parseable as a valid DOCX
        doc = Document(io.BytesIO(result))
        assert doc is not None

    def test_empty_content_generates_valid_doc(self):
        content = TORContent()
        result = self.generator.generate(content)
        doc = Document(io.BytesIO(result))
        assert doc is not None


class TestDOCXPageSetup:
    """Tests for page setup (margins, orientation, size)."""

    def setup_method(self):
        self.generator = DOCXGenerator()
        content = TORContent(project_name="Test")
        docx_bytes = self.generator.generate(content)
        self.doc = Document(io.BytesIO(docx_bytes))

    def test_margins_are_2_5_cm(self):
        section = self.doc.sections[0]
        # Allow small tolerance for rounding
        tolerance = Cm(0.01)
        assert abs(section.top_margin - PAGE_MARGIN) <= tolerance
        assert abs(section.bottom_margin - PAGE_MARGIN) <= tolerance
        assert abs(section.left_margin - PAGE_MARGIN) <= tolerance
        assert abs(section.right_margin - PAGE_MARGIN) <= tolerance

    def test_page_is_a4(self):
        section = self.doc.sections[0]
        # A4 = 21cm x 29.7cm
        assert abs(section.page_width - Cm(21.0)) <= Cm(0.1)
        assert abs(section.page_height - Cm(29.7)) <= Cm(0.1)


class TestDOCXFontSetup:
    """Tests for font configuration."""

    def setup_method(self):
        self.generator = DOCXGenerator()
        content = TORContent(project_name="Test")
        docx_bytes = self.generator.generate(content)
        self.doc = Document(io.BytesIO(docx_bytes))

    def test_normal_style_font_name(self):
        style = self.doc.styles["Normal"]
        assert style.font.name == FONT_NAME

    def test_normal_style_font_size(self):
        style = self.doc.styles["Normal"]
        assert style.font.size == BODY_FONT_SIZE


class TestDOCXContent:
    """Tests for document content rendering."""

    def setup_method(self):
        self.generator = DOCXGenerator()

    def test_title_is_present(self):
        content = TORContent(project_name="โครงการระบบ ICT")
        docx_bytes = self.generator.generate(content)
        doc = Document(io.BytesIO(docx_bytes))

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "ร่างขอบเขตของงาน" in full_text
        assert "Terms of Reference" in full_text

    def test_project_name_in_document(self):
        content = TORContent(project_name="โครงการพัฒนาระบบ")
        docx_bytes = self.generator.generate(content)
        doc = Document(io.BytesIO(docx_bytes))

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "โครงการพัฒนาระบบ" in full_text

    def test_all_13_sections_present(self):
        content = TORContent(project_name="Test")
        docx_bytes = self.generator.generate(content)
        doc = Document(io.BytesIO(docx_bytes))

        full_text = "\n".join(p.text for p in doc.paragraphs)
        for label in TOR_SECTION_LABELS.values():
            assert label in full_text, f"Section '{label}' not found in document"

    def test_section_content_rendered(self):
        content = TORContent(
            project_name="Test",
            sections={
                "s1": "ด้วยกระทรวงดิจิทัลมีความประสงค์จัดทำระบบ",
                "s2": "เพื่อพัฒนาระบบสารสนเทศ",
            },
        )
        docx_bytes = self.generator.generate(content)
        doc = Document(io.BytesIO(docx_bytes))

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "ด้วยกระทรวงดิจิทัลมีความประสงค์จัดทำระบบ" in full_text
        assert "เพื่อพัฒนาระบบสารสนเทศ" in full_text

    def test_empty_sections_show_placeholder(self):
        content = TORContent(project_name="Test", sections={})
        docx_bytes = self.generator.generate(content)
        doc = Document(io.BytesIO(docx_bytes))

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "(ยังไม่ได้กรอกข้อมูล)" in full_text

    def test_sub_sections_rendered(self):
        content = TORContent(
            project_name="Test",
            sections={"s4": "ขอบเขตของงานโดยรวม"},
            sub_sections={
                "s4": {
                    "4.1": "งานออกแบบระบบ",
                    "4.2": "งานพัฒนาระบบ",
                }
            },
        )
        docx_bytes = self.generator.generate(content)
        doc = Document(io.BytesIO(docx_bytes))

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "งานออกแบบระบบ" in full_text
        assert "งานพัฒนาระบบ" in full_text


class TestDOCXThaiNumerals:
    """Tests for Thai numeral section numbering."""

    def setup_method(self):
        self.generator = DOCXGenerator()

    def test_arabic_numerals_by_default(self):
        content = TORContent(project_name="Test", use_thai_numerals=False)
        docx_bytes = self.generator.generate(content)
        doc = Document(io.BytesIO(docx_bytes))

        full_text = "\n".join(p.text for p in doc.paragraphs)
        # Should have "1. ความเป็นมา"
        assert "1. ความเป็นมา" in full_text

    def test_thai_numerals_when_configured(self):
        content = TORContent(project_name="Test", use_thai_numerals=True)
        docx_bytes = self.generator.generate(content)
        doc = Document(io.BytesIO(docx_bytes))

        full_text = "\n".join(p.text for p in doc.paragraphs)
        # Should have "๑. ความเป็นมา"
        assert "๑. ความเป็นมา" in full_text


class TestDOCXDateFormatting:
    """Tests for Thai Buddhist Era date in document."""

    def setup_method(self):
        self.generator = DOCXGenerator()

    def test_date_in_buddhist_era(self):
        content = TORContent(
            project_name="Test",
            export_date=date(2024, 8, 15),
        )
        docx_bytes = self.generator.generate(content)
        doc = Document(io.BytesIO(docx_bytes))

        full_text = "\n".join(p.text for p in doc.paragraphs)
        # Buddhist Era: 2024 + 543 = 2567
        assert "2567" in full_text
        assert "สิงหาคม" in full_text

    def test_date_with_thai_numerals(self):
        content = TORContent(
            project_name="Test",
            export_date=date(2024, 8, 15),
            use_thai_numerals=True,
        )
        docx_bytes = self.generator.generate(content)
        doc = Document(io.BytesIO(docx_bytes))

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "๒๕๖๗" in full_text


class TestDOCXHeader:
    """Tests for document header (ministry name)."""

    def setup_method(self):
        self.generator = DOCXGenerator()

    def test_ministry_in_header(self):
        content = TORContent(
            project_name="Test",
            ministry="กระทรวงดิจิทัลเพื่อเศรษฐกิจและสังคม",
        )
        docx_bytes = self.generator.generate(content)
        doc = Document(io.BytesIO(docx_bytes))

        # Check header content
        header = doc.sections[0].header
        header_text = "".join(p.text for p in header.paragraphs)
        assert "กระทรวงดิจิทัลเพื่อเศรษฐกิจและสังคม" in header_text


class TestDOCXSectionOrder:
    """Tests for proper section ordering."""

    def test_section_order_is_complete(self):
        """All 13 sections must be in the order list."""
        assert len(TOR_SECTION_ORDER) == 13

    def test_all_sections_have_labels(self):
        """All ordered sections must have Thai labels."""
        for key in TOR_SECTION_ORDER:
            assert key in TOR_SECTION_LABELS

    def test_sections_appear_in_order(self):
        """Sections in the generated document appear in correct order."""
        generator = DOCXGenerator()
        content = TORContent(project_name="Test")
        docx_bytes = generator.generate(content)
        doc = Document(io.BytesIO(docx_bytes))

        full_text = "\n".join(p.text for p in doc.paragraphs)
        # Check that section 1 appears before section 2, etc.
        pos_s1 = full_text.find("ความเป็นมา")
        pos_s2 = full_text.find("วัตถุประสงค์")
        pos_s3 = full_text.find("คุณสมบัติของผู้เสนอราคา")

        assert pos_s1 < pos_s2 < pos_s3
