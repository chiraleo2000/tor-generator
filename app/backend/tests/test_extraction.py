"""Unit tests for document text extraction module.

Tests the text/DOCX extractors and the routing logic.
OCR functionality is tested with mocks since Tesseract may not be installed.
"""

from pathlib import Path
from unittest.mock import MagicMock, patch

import fitz  # PyMuPDF
import pytest
from docx import Document

from app.rag.extraction import (
    ExtractionResult,
    _deduplicate_adjacent,
    extract_combined_kb_json,
    extract_docx,
    extract_pdf,
    extract_text,
    extract_text_file,
    extract_tor_extract_json,
    flatten_decision_rules_json,
    ocr_page,
)

# =============================================================================
# Fixtures
# =============================================================================


@pytest.fixture
def sample_pdf_with_text(tmp_path: Path) -> str:
    """Create a sample PDF file with direct text content."""
    pdf_path = tmp_path / "sample.pdf"
    doc = fitz.open()

    # Add pages with Thai and English text
    page1 = doc.new_page()
    page1.insert_text(
        (72, 72),
        "This is the first page of the document. It contains English text.",
        fontsize=12,
    )

    page2 = doc.new_page()
    page2.insert_text(
        (72, 72),
        "Second page with more content for testing extraction.",
        fontsize=12,
    )

    doc.save(str(pdf_path))
    doc.close()
    return str(pdf_path)


@pytest.fixture
def sample_pdf_empty_page(tmp_path: Path) -> str:
    """Create a PDF with one text page and one empty page (simulating scanned)."""
    pdf_path = tmp_path / "mixed.pdf"
    doc = fitz.open()

    # Page 1: has text
    page1 = doc.new_page()
    page1.insert_text(
        (72, 72),
        "This page has text content that should be extracted directly.",
        fontsize=12,
    )

    # Page 2: empty (simulates scanned page)
    doc.new_page()

    doc.save(str(pdf_path))
    doc.close()
    return str(pdf_path)


@pytest.fixture
def sample_docx(tmp_path: Path) -> str:
    """Create a sample DOCX file with headings, paragraphs, and tables."""
    docx_path = tmp_path / "sample.docx"
    doc = Document()

    # Add heading
    doc.add_heading("Document Title", level=1)

    # Add paragraphs
    doc.add_paragraph("This is the first paragraph of the document.")
    doc.add_paragraph("This is the second paragraph with more details.")

    # Add sub-heading
    doc.add_heading("Section One", level=2)
    doc.add_paragraph("Content under section one.")

    # Add a table
    table = doc.add_table(rows=3, cols=3)
    headers = ["Header A", "Header B", "Header C"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    table.rows[1].cells[0].text = "Row 1 Col A"
    table.rows[1].cells[1].text = "Row 1 Col B"
    table.rows[1].cells[2].text = "Row 1 Col C"
    table.rows[2].cells[0].text = "Row 2 Col A"
    table.rows[2].cells[1].text = "Row 2 Col B"
    table.rows[2].cells[2].text = "Row 2 Col C"

    doc.save(str(docx_path))
    return str(docx_path)


@pytest.fixture
def sample_docx_thai(tmp_path: Path) -> str:
    """Create a DOCX file with Thai content."""
    docx_path = tmp_path / "thai.docx"
    doc = Document()

    doc.add_heading("ขอบเขตของงาน", level=1)
    doc.add_paragraph("รายละเอียดขอบเขตของงานจ้างที่ปรึกษา")
    doc.add_heading("วัตถุประสงค์", level=2)
    doc.add_paragraph("เพื่อศึกษาและวิเคราะห์ข้อมูลที่เกี่ยวข้อง")

    doc.save(str(docx_path))
    return str(docx_path)


@pytest.fixture
def sample_text_file(tmp_path: Path) -> str:
    """Create a plain text file."""
    txt_path = tmp_path / "sample.txt"
    txt_path.write_text(
        "This is a plain text file.\nIt has multiple lines.\nThird line here.",
        encoding="utf-8",
    )
    return str(txt_path)


@pytest.fixture
def sample_text_file_thai(tmp_path: Path) -> str:
    """Create a Thai plain text file in UTF-8."""
    txt_path = tmp_path / "thai.txt"
    txt_path.write_text(
        "นี่คือไฟล์ข้อความธรรมดา\nบรรทัดที่สอง\nบรรทัดที่สาม",
        encoding="utf-8",
    )
    return str(txt_path)


# =============================================================================
# Tests: extract_text (main entry point routing)
# =============================================================================


class TestExtractText:
    """Tests for the main extract_text routing function."""

    def test_routes_pdf(self, sample_pdf_with_text: str):
        """PDF MIME type routes to PDF extractor."""
        result = extract_text(sample_pdf_with_text, "application/pdf")
        assert isinstance(result, ExtractionResult)
        assert result.page_count == 2
        assert "first page" in result.text

    def test_routes_docx(self, sample_docx: str):
        """DOCX MIME type routes to DOCX extractor."""
        result = extract_text(
            sample_docx,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        assert isinstance(result, ExtractionResult)
        assert "Document Title" in result.text

    def test_routes_text(self, sample_text_file: str):
        """text/plain MIME type routes to text extractor."""
        result = extract_text(sample_text_file, "text/plain")
        assert isinstance(result, ExtractionResult)
        assert "plain text file" in result.text

    def test_routes_text_csv(self, tmp_path: Path):
        """text/* MIME subtypes also route to text extractor."""
        csv_path = tmp_path / "data.csv"
        csv_path.write_text("col1,col2\nval1,val2", encoding="utf-8")
        result = extract_text(str(csv_path), "text/csv")
        assert "col1,col2" in result.text

    def test_file_not_found(self):
        """Raises FileNotFoundError for missing files."""
        with pytest.raises(FileNotFoundError):
            extract_text("/nonexistent/file.pdf", "application/pdf")

    def test_unsupported_mime_type(self, sample_text_file: str):
        """Raises ValueError for unsupported MIME types."""
        with pytest.raises(ValueError, match="Unsupported MIME type"):
            extract_text(sample_text_file, "application/octet-stream")

    def test_mime_type_case_insensitive(self, sample_pdf_with_text: str):
        """MIME type matching is case-insensitive."""
        result = extract_text(sample_pdf_with_text, "Application/PDF")
        assert result.page_count == 2

    def test_mime_type_whitespace_stripped(self, sample_pdf_with_text: str):
        """MIME type with leading/trailing whitespace is handled."""
        result = extract_text(sample_pdf_with_text, "  application/pdf  ")
        assert result.page_count == 2


# =============================================================================
# Tests: extract_pdf
# =============================================================================


class TestExtractPdf:
    """Tests for PDF text extraction."""

    def test_extracts_text_from_text_pdf(self, sample_pdf_with_text: str):
        """Extracts text from a PDF with embedded text."""
        result = extract_pdf(sample_pdf_with_text)
        assert result.page_count == 2
        assert result.method == "direct"
        assert "first page" in result.text
        assert "Second page" in result.text
        assert result.warnings == []

    def test_method_is_direct_for_text_pdf(self, sample_pdf_with_text: str):
        """Method is 'direct' when all pages have text."""
        result = extract_pdf(sample_pdf_with_text)
        assert result.method == "direct"

    @patch("app.rag.extraction._ocr_pdf_page")
    def test_ocr_fallback_for_empty_pages(
        self, mock_ocr: MagicMock, sample_pdf_empty_page: str
    ):
        """Falls back to OCR for pages with insufficient text."""
        mock_ocr.return_value = "OCR extracted text from scanned page"

        result = extract_pdf(sample_pdf_empty_page)
        assert result.page_count == 2
        assert result.method == "mixed"
        assert "OCR extracted text" in result.text
        # OCR was called for the empty page
        mock_ocr.assert_called_once()

    @patch("app.rag.extraction._ocr_pdf_page")
    def test_ocr_failure_records_warning(
        self, mock_ocr: MagicMock, sample_pdf_empty_page: str
    ):
        """Records warning when OCR fails for a page."""
        mock_ocr.return_value = ""

        result = extract_pdf(sample_pdf_empty_page)
        assert len(result.warnings) >= 1
        assert any("Could not extract text" in w for w in result.warnings)

    def test_single_page_pdf(self, tmp_path: Path):
        """Handles single-page PDF correctly."""
        pdf_path = tmp_path / "single.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Single page content here.", fontsize=12)
        doc.save(str(pdf_path))
        doc.close()

        result = extract_pdf(str(pdf_path))
        assert result.page_count == 1
        assert "Single page content" in result.text

    @patch("app.rag.extraction.ocr_page", return_value="OCR from rendered PNG")
    def test_ocr_renders_temp_png_then_calls_tesseract(
        self, mock_ocr: MagicMock, sample_pdf_empty_page: str
    ):
        """Writes a temp PNG (closed first on Windows) then OCRs that path."""
        result = extract_pdf(sample_pdf_empty_page)
        assert "OCR from rendered PNG" in result.text
        mock_ocr.assert_called_once()
        image_path = Path(mock_ocr.call_args.args[0])
        assert image_path.suffix.lower() == ".png"

    @patch("app.rag.extraction._ocr_pdf_page")
    def test_all_pages_ocr(self, mock_ocr: MagicMock, tmp_path: Path):
        """When all pages need OCR, method is 'ocr'."""
        # Create PDF with no text (just empty pages)
        pdf_path = tmp_path / "scanned.pdf"
        doc = fitz.open()
        doc.new_page()
        doc.new_page()
        doc.save(str(pdf_path))
        doc.close()

        mock_ocr.return_value = "OCR text"

        result = extract_pdf(str(pdf_path))
        assert result.method == "ocr"
        assert mock_ocr.call_count == 2


# =============================================================================
# Tests: extract_docx
# =============================================================================


class TestExtractDocx:
    """Tests for DOCX text extraction."""

    def test_extracts_paragraphs(self, sample_docx: str):
        """Extracts paragraph text from DOCX."""
        result = extract_docx(sample_docx)
        assert "first paragraph" in result.text
        assert "second paragraph" in result.text

    def test_preserves_heading_structure(self, sample_docx: str):
        """Headings are marked with # prefix indicating level."""
        result = extract_docx(sample_docx)
        assert "# Document Title" in result.text
        assert "## Section One" in result.text

    def test_extracts_table_content(self, sample_docx: str):
        """Tables are extracted with cell content."""
        result = extract_docx(sample_docx)
        assert "Header A" in result.text
        assert "Row 1 Col A" in result.text
        assert "[Table 1]" in result.text

    def test_thai_content(self, sample_docx_thai: str):
        """Handles Thai content correctly."""
        result = extract_docx(sample_docx_thai)
        assert "ขอบเขตของงาน" in result.text
        assert "วัตถุประสงค์" in result.text

    def test_method_is_direct(self, sample_docx: str):
        """DOCX extraction method is always 'direct'."""
        result = extract_docx(sample_docx)
        assert result.method == "direct"

    def test_page_count_estimation(self, sample_docx: str):
        """Page count is estimated from content length."""
        result = extract_docx(sample_docx)
        assert result.page_count >= 1

    def test_empty_docx(self, tmp_path: Path):
        """Handles empty DOCX file."""
        docx_path = tmp_path / "empty.docx"
        doc = Document()
        doc.save(str(docx_path))

        result = extract_docx(str(docx_path))
        assert result.text == ""
        assert result.page_count == 1

    def test_invalid_docx_raises_error(self, tmp_path: Path):
        """Raises ValueError for invalid DOCX files."""
        bad_path = tmp_path / "bad.docx"
        bad_path.write_bytes(b"not a valid docx file")
        path = str(bad_path)

        with pytest.raises(ValueError, match="Failed to open DOCX"):
            extract_docx(path)


# =============================================================================
# Tests: extract_text_file
# =============================================================================


class TestExtractTextFile:
    """Tests for plain text file extraction."""

    def test_reads_utf8_file(self, sample_text_file: str):
        """Reads UTF-8 encoded text files."""
        result = extract_text_file(sample_text_file)
        assert "plain text file" in result.text
        assert "multiple lines" in result.text
        assert result.page_count == 1
        assert result.method == "direct"

    def test_reads_thai_utf8(self, sample_text_file_thai: str):
        """Reads Thai UTF-8 text correctly."""
        result = extract_text_file(sample_text_file_thai)
        assert "ไฟล์ข้อความธรรมดา" in result.text

    def test_strips_whitespace(self, tmp_path: Path):
        """Strips leading/trailing whitespace from content."""
        path = tmp_path / "spaces.txt"
        path.write_text("  content with spaces  \n\n", encoding="utf-8")

        result = extract_text_file(str(path))
        assert result.text == "content with spaces"

    def test_fallback_to_thai_encoding(self, tmp_path: Path):
        """Falls back to Windows-874 encoding for Thai files not in UTF-8."""
        path = tmp_path / "thai_874.txt"
        # Write Thai text in Windows-874 encoding
        thai_text = "สวัสดีครับ"
        path.write_bytes(thai_text.encode("cp874"))

        result = extract_text_file(str(path))
        assert "สวัสดีครับ" in result.text
        assert len(result.warnings) == 1
        assert "Windows-874" in result.warnings[0]

    def test_empty_file(self, tmp_path: Path):
        """Handles empty text files."""
        path = tmp_path / "empty.txt"
        path.write_text("", encoding="utf-8")

        result = extract_text_file(str(path))
        assert result.text == ""
        assert result.page_count == 1


# =============================================================================
# Tests: ocr_page
# =============================================================================


class TestOcrPage:
    """Tests for the OCR function (mocked since Tesseract may not be available)."""

    @patch("app.rag.extraction.subprocess.run")
    def test_successful_ocr(self, mock_run: MagicMock, tmp_path: Path):
        """Returns extracted text on successful OCR."""
        mock_run.return_value = MagicMock(
            returncode=0,
            stdout="Extracted Thai text สวัสดี\n",
            stderr="",
        )

        img_path = tmp_path / "page.png"
        img_path.write_bytes(b"\x89PNG\r\n")  # Minimal PNG header

        result = ocr_page(str(img_path), lang="tha+eng", timeout=30)
        assert result == "Extracted Thai text สวัสดี"

        # Verify tesseract was called correctly
        call_args = mock_run.call_args
        cmd = call_args[0][0]
        assert cmd[0] == "tesseract"
        assert "-l" in cmd
        assert "tha+eng" in cmd
        assert call_args[1]["timeout"] == 30
        assert call_args[1]["encoding"] == "utf-8"
        assert call_args[1]["errors"] == "replace"

    @patch("app.rag.extraction.subprocess.run")
    def test_ocr_timeout(self, mock_run: MagicMock, tmp_path: Path):
        """Raises TimeoutExpired when OCR exceeds timeout."""
        import subprocess

        mock_run.side_effect = subprocess.TimeoutExpired(cmd="tesseract", timeout=30)

        img_path = tmp_path / "page.png"
        img_path.write_bytes(b"\x89PNG\r\n")
        path = str(img_path)

        with pytest.raises(subprocess.TimeoutExpired):
            ocr_page(path, timeout=30)

    @patch("app.rag.extraction.subprocess.run")
    def test_tesseract_not_found(self, mock_run: MagicMock, tmp_path: Path):
        """Raises FileNotFoundError if Tesseract is not installed."""
        mock_run.side_effect = FileNotFoundError("tesseract not found")

        img_path = tmp_path / "page.png"
        img_path.write_bytes(b"\x89PNG\r\n")
        path = str(img_path)

        with pytest.raises(FileNotFoundError, match="Tesseract OCR is not installed"):
            ocr_page(path)

    @patch("app.rag.extraction.subprocess.run")
    def test_tesseract_error(self, mock_run: MagicMock, tmp_path: Path):
        """Raises RuntimeError if Tesseract exits with non-zero."""
        mock_run.return_value = MagicMock(
            returncode=1,
            stdout="",
            stderr="Error opening data file",
        )

        img_path = tmp_path / "page.png"
        img_path.write_bytes(b"\x89PNG\r\n")
        path = str(img_path)

        with pytest.raises(RuntimeError, match="Tesseract OCR failed"):
            ocr_page(path)


# =============================================================================
# Tests: _deduplicate_adjacent helper
# =============================================================================


class TestDeduplicateAdjacent:
    """Tests for the adjacent deduplication helper."""

    def test_no_duplicates(self):
        assert _deduplicate_adjacent(["a", "b", "c"]) == ["a", "b", "c"]

    def test_adjacent_duplicates(self):
        assert _deduplicate_adjacent(["a", "a", "b", "b", "c"]) == ["a", "b", "c"]

    def test_non_adjacent_duplicates_kept(self):
        assert _deduplicate_adjacent(["a", "b", "a"]) == ["a", "b", "a"]

    def test_empty_list(self):
        assert _deduplicate_adjacent([]) == []

    def test_single_item(self):
        assert _deduplicate_adjacent(["a"]) == ["a"]

    def test_all_same(self):
        assert _deduplicate_adjacent(["x", "x", "x"]) == ["x"]


# =============================================================================
# Tests: ExtractionResult dataclass
# =============================================================================


class TestExtractionResult:
    """Tests for ExtractionResult dataclass behavior."""

    def test_default_warnings_empty(self):
        """Warnings default to empty list."""
        result = ExtractionResult(text="hello", page_count=1, method="direct")
        assert result.warnings == []

    def test_with_warnings(self):
        """Can create with warnings."""
        result = ExtractionResult(
            text="partial",
            page_count=2,
            method="ocr",
            warnings=["OCR timeout on page 1"],
        )
        assert len(result.warnings) == 1
        assert "timeout" in result.warnings[0]


class TestExtractTorJson:
    """Tests for knowledge-base *_tor_extract.json parsing."""

    def test_extracts_focus_area_content(self, tmp_path: Path):
        path = tmp_path / "กฎกระทรวงวงเงิน_tor_extract.json"
        path.write_text(
            '{"source_file": "rule.txt", "focus_areas": {'
            '"procurement_methods": [{"content": "วิธีเฉพาะเจาะจง วงเงินไม่เกินห้าแสนบาท"}]}}',
            encoding="utf-8",
        )
        result = extract_tor_extract_json(str(path))
        assert "วิธีเฉพาะเจาะจง" in result.text
        assert "rule.txt" in result.text
        assert result.method == "direct"

    def test_empty_focus_areas_returns_empty_text(self, tmp_path: Path):
        path = tmp_path / "empty_tor_extract.json"
        path.write_text(
            '{"source_file": "x.txt", "focus_areas": {"definitions": []}}',
            encoding="utf-8",
        )
        result = extract_tor_extract_json(str(path))
        assert result.text == ""
        assert result.warnings

    def test_routes_json_mime_and_filename(self, tmp_path: Path):
        path = tmp_path / "พัสดุ_tor_extract.json"
        path.write_text(
            '{"focus_areas": {"definitions": [{"content": "พัสดุ หมายถึง"}]}}',
            encoding="utf-8",
        )
        via_mime = extract_text(str(path), "application/json")
        via_name = extract_text(str(path), "text/plain")
        assert "พัสดุ" in via_mime.text
        assert "พัสดุ" in via_name.text


class TestExtractCombinedKbJson:
    """Tests for knowledge-base *_combined.json parsing."""

    def test_extracts_sections_content(self, tmp_path: Path):
        path = tmp_path / "_definitions_combined.json"
        path.write_text(
            '{"name": "นิยามและประเภท", "sections": ['
            '{"content": "พัสดุ หมายถึง วัสดุ ครุภัณฑ์"}]}',
            encoding="utf-8",
        )
        result = extract_combined_kb_json(str(path))
        assert "พัสดุ หมายถึง" in result.text
        assert "นิยามและประเภท" in result.text
        assert "focus_areas" not in result.text
        assert result.method == "direct"

    def test_empty_sections_returns_empty_text(self, tmp_path: Path):
        path = tmp_path / "_empty_combined.json"
        path.write_text('{"sections": []}', encoding="utf-8")
        result = extract_combined_kb_json(str(path))
        assert result.text == ""
        assert result.warnings

    def test_routes_combined_filename(self, tmp_path: Path):
        path = tmp_path / "_procurement_methods_combined.json"
        path.write_text(
            '{"sections": [{"content": "วิธีเฉพาะเจาะจง วงเงินไม่เกินห้าแสนบาท"}]}',
            encoding="utf-8",
        )
        via_mime = extract_text(str(path), "application/json")
        via_name = extract_text(str(path), "text/plain")
        assert "วิธีเฉพาะเจาะจง" in via_mime.text
        assert "วิธีเฉพาะเจาะจง" in via_name.text


class TestFlattenDecisionRulesJson:
    """Tests for 04-decision-rules nested JSON flattening."""

    def test_flattens_nested_rules_to_thai(self, tmp_path: Path):
        rules_dir = tmp_path / "04-decision-rules"
        rules_dir.mkdir()
        path = rules_dir / "method_selection.json"
        path.write_text(
            '{"title": "การเลือกวิธีจัดซื้อจัดจ้าง", "rules": {'
            '"step_1_primary_method": {'
            '"description": "เลือกวิธีจัดซื้อจัดจ้างหลัก",'
            '"rules": [{'
            '"id": "R1",'
            '"condition": "วงเงินไม่เกิน 500,000 บาท",'
            '"result": "เฉพาะเจาะจง",'
            '"legal_basis": "มาตรา 56(2)(ข)"'
            "}]}}}",
            encoding="utf-8",
        )
        result = flatten_decision_rules_json(str(path))
        assert "เลือกวิธีจัดซื้อจัดจ้างหลัก" in result.text
        assert "วงเงินไม่เกิน 500,000 บาท" in result.text
        assert "เฉพาะเจาะจง" in result.text
        assert "step_1_primary_method" not in result.text
        assert "R1" not in result.text
        assert result.method == "direct"

    def test_empty_rules_returns_empty_text(self, tmp_path: Path):
        path = tmp_path / "empty_rules.json"
        path.write_text("{}", encoding="utf-8")
        result = flatten_decision_rules_json(str(path))
        assert result.text == ""
        assert result.warnings

    def test_routes_decision_rules_path(self, tmp_path: Path):
        rules_dir = tmp_path / "04-decision-rules"
        rules_dir.mkdir()
        path = rules_dir / "method_selection.json"
        path.write_text(
            '{"rules": {"step_1_primary_method": {'
            '"description": "เลือกวิธีประกาศเชิญชวนทั่วไป"}}}',
            encoding="utf-8",
        )
        via_mime = extract_text(str(path), "application/json")
        via_name = extract_text(str(path), "text/plain")
        assert "เลือกวิธีประกาศเชิญชวนทั่วไป" in via_mime.text
        assert "เลือกวิธีประกาศเชิญชวนทั่วไป" in via_name.text
        assert "step_1_primary_method" not in via_mime.text

