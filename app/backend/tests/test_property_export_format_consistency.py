"""Property-based tests for Export Format Consistency (Property 10).

Verifies that DOCX and PDF generators produce documents with identical textual
content when given the same TORContent input. Since extracting text from actual
PDF bytes requires external tools, we compare the PDF's intermediate HTML
representation against the DOCX paragraphs for content equivalence.

**Validates: Requirements 8.1, 8.2**

# Feature: tor-drafting-review-app, Property 10: Export Format Consistency
"""

from __future__ import annotations

import io
import re
from datetime import date
from html.parser import HTMLParser

import pytest
from docx import Document
from hypothesis import given, settings
from hypothesis import strategies as st

from app.export.docx_generator import (
    TOR_SECTION_LABELS,
    TOR_SECTION_ORDER,
    DOCXGenerator,
    TORContent,
)
from app.export.pdf_generator import PDFGenerator


# ---------------------------------------------------------------------------
# Helpers: text extraction
# ---------------------------------------------------------------------------


class _HTMLTextExtractor(HTMLParser):
    """Extract visible text content from HTML, ignoring tags and styles."""

    def __init__(self) -> None:
        super().__init__()
        self._text_parts: list[str] = []
        self._skip_tags = {"style", "script", "head"}
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in self._skip_tags:
            self._skip_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in self._skip_tags and self._skip_depth > 0:
            self._skip_depth -= 1

    def handle_data(self, data: str) -> None:
        if self._skip_depth == 0:
            stripped = data.strip()
            if stripped:
                self._text_parts.append(stripped)

    def get_text_segments(self) -> list[str]:
        """Return the list of non-empty text segments found."""
        return self._text_parts


def _extract_text_from_html(html: str) -> set[str]:
    """Extract all visible text segments from HTML as a set of non-empty strings."""
    extractor = _HTMLTextExtractor()
    extractor.feed(html)
    return set(extractor.get_text_segments())


def _extract_text_from_docx_bytes(docx_bytes: bytes) -> set[str]:
    """Extract all paragraph text from DOCX bytes as a set of non-empty strings."""
    doc = Document(io.BytesIO(docx_bytes))
    text_parts: set[str] = set()

    # Body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_parts.add(text)

    # Header text
    for section in doc.sections:
        for para in section.header.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.add(text)

        # Footer may contain page number fields (non-textual), skip those

    return text_parts


def _normalize_text(text: str) -> str:
    """Normalize whitespace in text for comparison."""
    return re.sub(r"\s+", " ", text).strip()


# ---------------------------------------------------------------------------
# Strategies
# ---------------------------------------------------------------------------

# Thai-like text for realistic section content
_thai_text_chars = st.text(
    alphabet=st.sampled_from(
        "กขคงจฉชซฌญฎฏฐฑฒณดตถทธนบปผฝพฟภมยรลวศษสหฬอฮ"
        "ะาิีึืุูเแโใไๆ็่้๊๋์"
        " 0123456789"
    ),
    min_size=1,
    max_size=100,
)

# Project name: non-empty text
_project_name_strategy = st.text(
    alphabet=st.sampled_from(
        "กขคงจฉชซญดตถทธนบปพฟมยรลวศษสหอฮ"
        "ะาิีึืุูเแโใไ"
        " "
        "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz"
        "0123456789"
    ),
    min_size=1,
    max_size=60,
).filter(lambda s: s.strip() != "")

# Ministry name
_ministry_strategy = st.text(
    alphabet=st.sampled_from(
        "กขคงจฉชซญดตถทธนบปพฟมยรลวศษสหอฮ"
        "ะาิีึืุูเแโใไ"
        " "
    ),
    min_size=1,
    max_size=50,
).filter(lambda s: s.strip() != "")

# Section content: paragraphs separated by newlines, non-empty
_section_content_strategy = st.text(
    alphabet=st.sampled_from(
        "กขคงจฉชซญดตถทธนบปพฟมยรลวศษสหอฮ"
        "ะาิีึืุูเแโใไ"
        " 0123456789.\n"
    ),
    min_size=1,
    max_size=200,
).filter(lambda s: any(c.strip() for c in s.split("\n") if c.strip()))

# Sub-section key like "4.1", "4.2"
_sub_key_strategy = st.builds(
    lambda parent, child: f"{parent}.{child}",
    st.integers(min_value=1, max_value=13),
    st.integers(min_value=1, max_value=14),
)


@st.composite
def _tor_content_strategy(draw):
    """Generate a random TORContent with at least some sections filled."""
    project_name = draw(_project_name_strategy)
    ministry = draw(_ministry_strategy)
    budget = draw(st.integers(min_value=100_000, max_value=10_000_000_000))
    project_type = draw(st.sampled_from(["it", "construction", "consulting", "general"]))
    use_thai_numerals = draw(st.booleans())

    # Randomly fill some sections (at least 1, up to all 13)
    num_sections_to_fill = draw(st.integers(min_value=1, max_value=13))
    section_keys_to_fill = draw(
        st.lists(
            st.sampled_from(TOR_SECTION_ORDER),
            min_size=num_sections_to_fill,
            max_size=num_sections_to_fill,
            unique=True,
        )
    )

    sections: dict[str, str] = {}
    for key in section_keys_to_fill:
        content = draw(_section_content_strategy)
        sections[key] = content

    # Optionally add sub-sections for section s4 (scope of work)
    sub_sections: dict[str, dict[str, str]] = {}
    add_sub = draw(st.booleans())
    if add_sub and "s4" in sections:
        num_subs = draw(st.integers(min_value=1, max_value=5))
        subs: dict[str, str] = {}
        for i in range(1, num_subs + 1):
            sub_key = f"4.{i}"
            sub_content = draw(_section_content_strategy)
            subs[sub_key] = sub_content
        sub_sections["s4"] = subs

    # Use a fixed date for deterministic comparison
    export_date = date(2024, 8, 15)

    return TORContent(
        project_name=project_name,
        ministry=ministry,
        budget=budget,
        project_type=project_type,
        sections=sections,
        sub_sections=sub_sections,
        export_date=export_date,
        use_thai_numerals=use_thai_numerals,
    )


# ---------------------------------------------------------------------------
# Property Tests
# ---------------------------------------------------------------------------


@pytest.mark.property
class TestExportFormatConsistency:
    """Property 10: Export Format Consistency.

    For any TOR content, the DOCX and PDF generators produce documents
    containing the same textual content (section names, section content,
    project name, dates). We compare the PDF's intermediate HTML against
    the DOCX paragraph text.
    """

    @given(content=_tor_content_strategy())
    @settings(max_examples=100, deadline=None)
    # Feature: tor-drafting-review-app, Property 10: Export Format Consistency
    def test_docx_and_pdf_contain_same_section_headings(self, content: TORContent):
        """For any TOR content, both formats contain the same section headings.

        **Validates: Requirements 8.1, 8.2**
        """
        # Generate DOCX
        docx_gen = DOCXGenerator()
        docx_bytes = docx_gen.generate(content)
        docx_texts = _extract_text_from_docx_bytes(docx_bytes)

        # Get PDF HTML (intermediate representation)
        pdf_gen = PDFGenerator()
        html = pdf_gen._build_html(content)
        html_texts = _extract_text_from_html(html)

        # All 13 section headings should be present in both outputs
        from app.export.thai_formatting import format_section_number

        for idx, section_key in enumerate(TOR_SECTION_ORDER, start=1):
            num_str = format_section_number(idx, content.use_thai_numerals)
            label = TOR_SECTION_LABELS[section_key]
            heading_text = f"{num_str}. {label}"

            assert heading_text in docx_texts, (
                f"Section heading '{heading_text}' missing from DOCX"
            )
            assert heading_text in html_texts, (
                f"Section heading '{heading_text}' missing from PDF HTML"
            )

    @given(content=_tor_content_strategy())
    @settings(max_examples=100, deadline=None)
    # Feature: tor-drafting-review-app, Property 10: Export Format Consistency
    def test_docx_and_pdf_contain_same_project_name(self, content: TORContent):
        """For any TOR content with a project name, both formats contain that project name.

        **Validates: Requirements 8.1, 8.2**
        """
        docx_gen = DOCXGenerator()
        docx_bytes = docx_gen.generate(content)
        docx_texts = _extract_text_from_docx_bytes(docx_bytes)

        pdf_gen = PDFGenerator()
        html = pdf_gen._build_html(content)
        html_texts = _extract_text_from_html(html)

        # Project name should appear in both
        if content.project_name and content.project_name.strip():
            project_name_stripped = content.project_name.strip()
            assert project_name_stripped in docx_texts, (
                f"Project name '{project_name_stripped}' missing from DOCX"
            )
            assert project_name_stripped in html_texts, (
                f"Project name '{project_name_stripped}' missing from PDF HTML"
            )

    @given(content=_tor_content_strategy())
    @settings(max_examples=100, deadline=None)
    # Feature: tor-drafting-review-app, Property 10: Export Format Consistency
    def test_docx_and_pdf_contain_same_section_content(self, content: TORContent):
        """For any TOR content, both formats contain the same section body content.

        Each non-empty paragraph line within a section should appear in both outputs.

        **Validates: Requirements 8.1, 8.2**
        """
        docx_gen = DOCXGenerator()
        docx_bytes = docx_gen.generate(content)
        docx_texts = _extract_text_from_docx_bytes(docx_bytes)

        pdf_gen = PDFGenerator()
        html = pdf_gen._build_html(content)
        html_texts = _extract_text_from_html(html)

        # Every non-empty paragraph in sections should be in both
        for section_key, section_text in content.sections.items():
            paragraphs = section_text.strip().split("\n")
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                assert para in docx_texts, (
                    f"Section {section_key} paragraph '{para[:50]}...' missing from DOCX"
                )
                assert para in html_texts, (
                    f"Section {section_key} paragraph '{para[:50]}...' missing from PDF HTML"
                )

    @given(content=_tor_content_strategy())
    @settings(max_examples=100, deadline=None)
    # Feature: tor-drafting-review-app, Property 10: Export Format Consistency
    def test_docx_and_pdf_contain_same_date(self, content: TORContent):
        """For any TOR content, both formats contain the same formatted date.

        **Validates: Requirements 8.1, 8.2**
        """
        from app.export.thai_formatting import format_thai_date

        docx_gen = DOCXGenerator()
        docx_bytes = docx_gen.generate(content)
        docx_texts = _extract_text_from_docx_bytes(docx_bytes)

        pdf_gen = PDFGenerator()
        html = pdf_gen._build_html(content)
        html_texts = _extract_text_from_html(html)

        # The formatted date should be identical in both
        export_date = content.export_date or date.today()
        expected_date = format_thai_date(
            export_date, use_thai_numerals=content.use_thai_numerals
        )

        assert expected_date in docx_texts, (
            f"Date '{expected_date}' missing from DOCX"
        )
        assert expected_date in html_texts, (
            f"Date '{expected_date}' missing from PDF HTML"
        )

    @given(content=_tor_content_strategy())
    @settings(max_examples=100, deadline=None)
    # Feature: tor-drafting-review-app, Property 10: Export Format Consistency
    def test_docx_and_pdf_contain_same_sub_section_content(self, content: TORContent):
        """For any TOR content with sub-sections, both formats include sub-section text.

        **Validates: Requirements 8.1, 8.2**
        """
        docx_gen = DOCXGenerator()
        docx_bytes = docx_gen.generate(content)
        docx_texts = _extract_text_from_docx_bytes(docx_bytes)

        pdf_gen = PDFGenerator()
        html = pdf_gen._build_html(content)
        html_texts = _extract_text_from_html(html)

        for section_key, subs in content.sub_sections.items():
            for sub_key, sub_content in subs.items():
                if not sub_content or not sub_content.strip():
                    continue
                paragraphs = sub_content.strip().split("\n")
                for para in paragraphs:
                    para = para.strip()
                    if not para:
                        continue
                    assert para in docx_texts, (
                        f"Sub-section {sub_key} paragraph '{para[:50]}...' "
                        f"missing from DOCX"
                    )
                    assert para in html_texts, (
                        f"Sub-section {sub_key} paragraph '{para[:50]}...' "
                        f"missing from PDF HTML"
                    )

    @given(content=_tor_content_strategy())
    @settings(max_examples=100, deadline=None)
    # Feature: tor-drafting-review-app, Property 10: Export Format Consistency
    def test_docx_and_pdf_section_order_matches(self, content: TORContent):
        """For any TOR content, sections appear in the same order in both formats.

        **Validates: Requirements 8.1, 8.2**
        """
        from app.export.thai_formatting import format_section_number

        # PDF HTML preserves insertion order
        pdf_gen = PDFGenerator()
        html = pdf_gen._build_html(content)

        # Extract section headings from HTML in order
        heading_pattern = re.compile(
            r'<div class="section-heading">(.+?)</div>'
        )
        html_headings = [
            m.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
            .replace("&quot;", '"').replace("&#x27;", "'")
            for m in heading_pattern.findall(html)
        ]

        # DOCX paragraphs in order - extract bold paragraphs that look like section headings
        # Section headings match "N. <Thai label>" (not sub-section headings like "4.1")
        docx_gen = DOCXGenerator()
        docx_bytes = docx_gen.generate(content)
        doc = Document(io.BytesIO(docx_bytes))

        docx_headings: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Section headings are bold and match pattern "N. <label>" where label
            # contains Thai characters (distinguishes from sub-headings like "4.1")
            runs = para.runs
            if runs and runs[0].bold and re.match(
                r"^[\d๐-๙]+\.\s+\S", text
            ):
                docx_headings.append(text)

        # Both should have exactly 13 section headings in the same order
        assert len(html_headings) == 13, (
            f"Expected 13 section headings in HTML, got {len(html_headings)}"
        )
        assert len(docx_headings) == 13, (
            f"Expected 13 section headings in DOCX, got {len(docx_headings)}"
        )
        assert html_headings == docx_headings, (
            f"Section heading order differs:\n"
            f"  HTML: {html_headings[:3]}...\n"
            f"  DOCX: {docx_headings[:3]}..."
        )

    @given(content=_tor_content_strategy())
    @settings(max_examples=100, deadline=None)
    # Feature: tor-drafting-review-app, Property 10: Export Format Consistency
    def test_docx_and_pdf_ministry_header_present(self, content: TORContent):
        """For any TOR content with a ministry name, both formats include it.

        **Validates: Requirements 8.1, 8.2**
        """
        docx_gen = DOCXGenerator()
        docx_bytes = docx_gen.generate(content)
        docx_texts = _extract_text_from_docx_bytes(docx_bytes)

        pdf_gen = PDFGenerator()
        html = pdf_gen._build_html(content)
        html_texts = _extract_text_from_html(html)

        if content.ministry and content.ministry.strip():
            # Ministry should appear in both (as header)
            # Note: DOCX extraction strips text, so compare stripped ministry
            ministry_stripped = content.ministry.strip()
            assert ministry_stripped in docx_texts, (
                f"Ministry '{ministry_stripped}' missing from DOCX header"
            )
            assert ministry_stripped in html_texts, (
                f"Ministry '{ministry_stripped}' missing from PDF HTML header"
            )

    @given(content=_tor_content_strategy())
    @settings(max_examples=100, deadline=None)
    # Feature: tor-drafting-review-app, Property 10: Export Format Consistency
    def test_all_docx_content_text_found_in_pdf_html(self, content: TORContent):
        """For any TOR content, every meaningful text segment in DOCX also exists in PDF HTML.

        This is the comprehensive equivalence check: the set of textual content
        in DOCX should be a subset of the PDF HTML text (and vice versa for
        user-provided content).

        **Validates: Requirements 8.1, 8.2**
        """
        docx_gen = DOCXGenerator()
        docx_bytes = docx_gen.generate(content)
        docx_texts = _extract_text_from_docx_bytes(docx_bytes)

        pdf_gen = PDFGenerator()
        html = pdf_gen._build_html(content)
        html_texts = _extract_text_from_html(html)

        # Filter out purely structural text (separators, placeholders) for this check
        separator = "─" * 60
        placeholder = "(ยังไม่ได้กรอกข้อมูล)"
        title = "ร่างขอบเขตของงาน (Terms of Reference: TOR)"

        # The meaningful content that must be in both
        # (excluding separator formatting differences and title)
        meaningful_docx = {
            t for t in docx_texts
            if t != separator and t != placeholder and t != title
        }
        meaningful_html = {
            t for t in html_texts
            if t != separator and t != placeholder and t != title
        }

        # Every piece of user content in DOCX should be in PDF HTML
        missing_from_html = meaningful_docx - meaningful_html
        # Filter out page number field artifacts from DOCX
        missing_from_html = {
            t for t in missing_from_html
            if not t.startswith("PAGE")
        }

        assert missing_from_html == set(), (
            f"Content in DOCX not found in PDF HTML: "
            f"{list(missing_from_html)[:5]}"
        )

        # Every piece of user content in PDF HTML should be in DOCX
        missing_from_docx = meaningful_html - meaningful_docx
        assert missing_from_docx == set(), (
            f"Content in PDF HTML not found in DOCX: "
            f"{list(missing_from_docx)[:5]}"
        )
