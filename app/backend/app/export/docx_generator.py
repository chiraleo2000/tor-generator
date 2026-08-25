"""DOCX document generator for TOR export.

Generates Word (.docx) documents with Thai government formatting using python-docx.
Applies TH Sarabun New font, proper margins, headers, page numbering, and supports
Thai/Arabic numeral configuration for section numbering.

Requirements: 8.1, 8.3, 8.4, 16.3
"""

import io
from dataclasses import dataclass, field
from datetime import date, datetime

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.domain.tor_sections import TOR_SECTION_LABELS, TOR_SECTION_ORDER
from app.export.thai_formatting import (
    format_section_number,
    format_thai_date,
)

# Font constants
FONT_NAME = "TH Sarabun New"
FONT_NAME_FALLBACK = "TH SarabunPSK"
BODY_FONT_SIZE = Pt(14)
HEADING_FONT_SIZE = Pt(16)
HEADER_FONT_SIZE = Pt(12)

# Page margin (2.5 cm all around)
PAGE_MARGIN = Cm(2.5)
W_RFONTS = "w:rFonts"


@dataclass
class TORContent:
    """Structured TOR content for DOCX generation.

    Attributes:
        project_name: Name of the project.
        ministry: Ministry/organization name.
        budget: Total budget in baht.
        project_type: Project type (it, construction, consulting, general).
        sections: Dict mapping section_key to content text.
        sub_sections: Dict mapping section_key to sub-section dict (sub_key → content).
        export_date: Date to display on the document (defaults to today).
        use_thai_numerals: Whether to use Thai numerals for section numbering.
    """

    project_name: str = ""
    ministry: str = ""
    budget: int = 0
    project_type: str = "general"
    sections: dict[str, str] = field(default_factory=dict)
    sub_sections: dict[str, dict[str, str]] = field(default_factory=dict)
    export_date: date | datetime | None = None
    use_thai_numerals: bool = False


class DOCXGenerator:
    """Generates DOCX documents with Thai government formatting.

    Usage:
        generator = DOCXGenerator()
        content = TORContent(
            project_name="ระบบ ICT",
            ministry="กระทรวงดิจิทัลฯ",
            budget=5_000_000,
            sections={"s1": "ความเป็นมาของโครงการ..."},
        )
        docx_bytes = generator.generate(content)
    """

    def __init__(self) -> None:
        """Initialize the generator."""
        pass

    def generate(self, content: TORContent) -> bytes:
        """Generate a DOCX file from TOR content.

        Args:
            content: Structured TOR content to render.

        Returns:
            Bytes representing the generated .docx file.
        """
        doc = Document()

        self._configure_page_setup(doc)
        self._set_default_font(doc)
        self._add_header(doc, content)
        self._add_page_numbers(doc)
        self._add_title(doc, content)
        self._add_document_info(doc, content)
        self._add_sections(doc, content)

        # Write to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _configure_page_setup(self, doc: Document) -> None:
        """Configure page orientation, margins, and size (A4)."""
        section = doc.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT

        # A4 size
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

        # 2.5cm margins on all sides
        section.top_margin = PAGE_MARGIN
        section.bottom_margin = PAGE_MARGIN
        section.left_margin = PAGE_MARGIN
        section.right_margin = PAGE_MARGIN

    def _set_default_font(self, doc: Document) -> None:
        """Set the default document font to TH Sarabun New 14pt."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = FONT_NAME
        font.size = BODY_FONT_SIZE

        # Set the East Asian / complex script font for Thai support
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn(W_RFONTS))
        if rfonts is None:
            rfonts = OxmlElement(W_RFONTS)
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), FONT_NAME)
        rfonts.set(qn("w:hAnsi"), FONT_NAME)
        rfonts.set(qn("w:cs"), FONT_NAME)
        rfonts.set(qn("w:eastAsia"), FONT_NAME)

        # Also configure heading styles
        for heading_level in range(1, 4):
            style_name = f"Heading {heading_level}"
            if style_name in doc.styles:
                h_style = doc.styles[style_name]
                h_font = h_style.font
                h_font.name = FONT_NAME
                h_font.size = HEADING_FONT_SIZE
                h_font.bold = True
                h_font.color.rgb = RGBColor(0, 0, 0)

                h_rpr = h_style.element.get_or_add_rPr()
                h_rfonts = h_rpr.find(qn(W_RFONTS))
                if h_rfonts is None:
                    h_rfonts = OxmlElement(W_RFONTS)
                    h_rpr.append(h_rfonts)
                h_rfonts.set(qn("w:ascii"), FONT_NAME)
                h_rfonts.set(qn("w:hAnsi"), FONT_NAME)
                h_rfonts.set(qn("w:cs"), FONT_NAME)
                h_rfonts.set(qn("w:eastAsia"), FONT_NAME)

    def _add_header(self, doc: Document, content: TORContent) -> None:
        """Add document header with ministry/organization name."""
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = header_para.add_run(content.ministry or "")
        run.font.name = FONT_NAME
        run.font.size = HEADER_FONT_SIZE
        self._set_run_font_cs(run)

    def _add_page_numbers(self, doc: Document) -> None:
        """Add page numbers to the document footer (centered)."""
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page number field code
        run = footer_para.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_char_begin)

        run2 = footer_para.add_run()
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "
        run2._r.append(instr_text)

        run3 = footer_para.add_run()
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run3._r.append(fld_char_end)

    def _add_title(self, doc: Document, content: TORContent) -> None:
        """Add the document title (ร่างขอบเขตของงาน / TOR)."""
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.space_before = Pt(0)
        title_para.space_after = Pt(12)

        run = title_para.add_run("ร่างขอบเขตของงาน (Terms of Reference: TOR)")
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = HEADING_FONT_SIZE
        self._set_run_font_cs(run)

        # Project name subtitle
        if content.project_name:
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_para.space_after = Pt(6)
            run = subtitle_para.add_run(content.project_name)
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = HEADING_FONT_SIZE
            self._set_run_font_cs(run)

    def _add_document_info(self, doc: Document, content: TORContent) -> None:
        """Add document metadata (date, budget summary)."""
        export_date = content.export_date or date.today()
        date_str = format_thai_date(export_date, use_thai_numerals=content.use_thai_numerals)

        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        info_para.space_after = Pt(12)
        run = info_para.add_run(date_str)
        run.font.name = FONT_NAME
        run.font.size = BODY_FONT_SIZE
        self._set_run_font_cs(run)

        # Separator line
        doc.add_paragraph("─" * 60)

    def _add_sections(self, doc: Document, content: TORContent) -> None:
        """Add all TOR sections in order."""
        section_num = 1
        for section_key in TOR_SECTION_ORDER:
            section_content = content.sections.get(section_key, "")
            section_label = TOR_SECTION_LABELS.get(section_key, f"ส่วนที่ {section_key}")

            # Section heading
            num_str = format_section_number(section_num, content.use_thai_numerals)
            heading_text = f"{num_str}. {section_label}"

            heading_para = doc.add_paragraph()
            heading_para.space_before = Pt(12)
            heading_para.space_after = Pt(6)
            run = heading_para.add_run(heading_text)
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = HEADING_FONT_SIZE
            self._set_run_font_cs(run)

            # Section body content
            if section_content:
                self._add_section_content(doc, section_content, content.use_thai_numerals)
            else:
                # Placeholder for empty sections
                placeholder_para = doc.add_paragraph()
                placeholder_run = placeholder_para.add_run("(ยังไม่ได้กรอกข้อมูล)")
                placeholder_run.font.name = FONT_NAME
                placeholder_run.font.size = BODY_FONT_SIZE
                placeholder_run.italic = True
                placeholder_run.font.color.rgb = RGBColor(128, 128, 128)
                self._set_run_font_cs(placeholder_run)

            # Sub-sections (e.g., s4 → 4.1, 4.2, ... 4.14)
            sub_sections = content.sub_sections.get(section_key, {})
            if sub_sections:
                self._add_sub_sections(
                    doc, section_num, sub_sections, content.use_thai_numerals
                )

            section_num += 1

    def _add_section_content(
        self, doc: Document, text: str, _use_thai_numerals: bool
    ) -> None:
        """Add section body text, splitting by paragraphs."""
        paragraphs = text.strip().split("\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Cm(1.25)
            run = para.add_run(para_text)
            run.font.name = FONT_NAME
            run.font.size = BODY_FONT_SIZE
            self._set_run_font_cs(run)

    def _add_sub_sections(
        self,
        doc: Document,
        _parent_num: int,
        sub_sections: dict[str, str],
        use_thai_numerals: bool,
    ) -> None:
        """Add sub-sections (e.g. 4.1, 4.2, ...)."""
        # Sort sub-sections by their numeric key
        sorted_keys = sorted(
            sub_sections.keys(),
            key=lambda k: self._parse_sub_key(k),
        )

        for sub_key in sorted_keys:
            sub_content = sub_sections[sub_key]

            # Format sub-section number
            sub_num_str = format_section_number(sub_key, use_thai_numerals)

            sub_heading_para = doc.add_paragraph()
            sub_heading_para.space_before = Pt(6)
            sub_heading_para.space_after = Pt(3)
            sub_heading_para.paragraph_format.left_indent = Cm(1.0)

            run = sub_heading_para.add_run(f"{sub_num_str}")
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = BODY_FONT_SIZE
            self._set_run_font_cs(run)

            # Sub-section content
            if sub_content:
                paragraphs = sub_content.strip().split("\n")
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if not para_text:
                        continue
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Cm(1.5)
                    run = para.add_run(para_text)
                    run.font.name = FONT_NAME
                    run.font.size = BODY_FONT_SIZE
                    self._set_run_font_cs(run)

    def _parse_sub_key(self, key: str) -> tuple[int, ...]:
        """Parse a sub-section key like '4.1' into a sortable tuple (4, 1)."""
        parts = key.replace(".", " ").split()
        result = []
        for part in parts:
            try:
                result.append(int(part))
            except ValueError:
                result.append(0)
        return tuple(result)

    def _set_run_font_cs(self, run) -> None:
        """Set the complex script (cs) font on a run for proper Thai rendering."""
        rpr = run._r.get_or_add_rPr()
        rfonts = rpr.find(qn(W_RFONTS))
        if rfonts is None:
            rfonts = OxmlElement(W_RFONTS)
            rpr.append(rfonts)
        rfonts.set(qn("w:cs"), FONT_NAME)

        # Also set cs size to match
        sz_cs = rpr.find(qn("w:szCs"))
        if sz_cs is None:
            sz_cs = OxmlElement("w:szCs")
            rpr.append(sz_cs)
        # python-docx uses half-points for font size
        if run.font.size:
            sz_cs.set(qn("w:val"), str(int(run.font.size.pt * 2)))
