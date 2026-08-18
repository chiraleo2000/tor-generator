"""Document text extraction module.

Extracts text content from PDF, DOCX, and plain text files.
Provides OCR fallback for scanned PDFs using Tesseract with Thai + English support.

Requirements: 3.1, 14.1, 14.2, 14.3, 14.5
"""

from __future__ import annotations

import json
import logging
import os
import subprocess
import tempfile
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Literal

import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)

# Minimum character count per page to consider text extraction successful
MIN_TEXT_LENGTH = 10

# Default OCR timeout in seconds
DEFAULT_OCR_TIMEOUT = 30


class ExtractionMethod(str, Enum):
    """Method used for text extraction."""

    DIRECT = "direct"
    OCR = "ocr"
    MIXED = "mixed"  # Some pages direct, some OCR


@dataclass
class ExtractionResult:
    """Result of a document text extraction operation.

    Attributes:
        text: The extracted text content.
        page_count: Number of pages in the document (1 for non-paginated formats).
        method: The extraction method used (direct, ocr, or mixed).
        warnings: List of warning messages encountered during extraction.
    """

    text: str
    page_count: int
    method: Literal["direct", "ocr", "mixed"]
    warnings: list[str] = field(default_factory=list)


def extract_text(file_path: str, mime_type: str) -> ExtractionResult:
    """Main entry point for document text extraction.

    Routes to the appropriate extractor based on MIME type.

    Args:
        file_path: Path to the file to extract text from.
        mime_type: MIME type of the file (e.g. application/pdf, DOCX, text/plain).

    Returns:
        ExtractionResult with extracted text and metadata.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the MIME type is unsupported.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    mime_type = mime_type.lower().strip()

    if _is_kb_json(path, mime_type):
        return _extract_kb_json(file_path)

    if mime_type == "application/pdf":
        return extract_pdf(file_path)
    elif mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return extract_docx(file_path)
    elif mime_type.startswith("text/"):
        return extract_text_file(file_path)
    elif mime_type in ("image/png", "image/jpeg", "image/gif", "image/webp"):
        warnings: list[str] = []
        try:
            text = ocr_page(file_path)
        except FileNotFoundError:
            warnings.append("ไม่พบ Tesseract สำหรับ OCR รูปสแกน")
            text = ""
        except (RuntimeError, subprocess.TimeoutExpired) as exc:
            warnings.append(f"OCR รูปไม่สำเร็จ: {exc}")
            text = ""
        return ExtractionResult(
            text=text.strip(),
            page_count=1,
            method="ocr",
            warnings=warnings,
        )
    else:
        raise ValueError(
            f"Unsupported MIME type: {mime_type}. "
            "Supported types: application/pdf, application/json, "
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document, "
            "image/png, image/jpeg, text/*"
        )


def extract_pdf(file_path: str, ocr_timeout: int = DEFAULT_OCR_TIMEOUT) -> ExtractionResult:
    """Extract text from a PDF file using PyMuPDF.

    For each page, attempts direct text extraction first.
    If text is too short (< 10 chars), falls back to OCR for that page.

    Args:
        file_path: Path to the PDF file.
        ocr_timeout: Timeout in seconds for OCR per page.

    Returns:
        ExtractionResult with extracted text and metadata.
    """
    warnings: list[str] = []
    page_texts: list[str] = []
    used_ocr = False
    used_direct = False

    doc = fitz.open(file_path)
    page_count = len(doc)

    for page_num in range(page_count):
        page = doc[page_num]
        text = page.get_text().strip()

        if len(text) >= MIN_TEXT_LENGTH:
            # Direct text extraction succeeded
            page_texts.append(text)
            used_direct = True
        else:
            # Text is too short/empty — fall back to OCR
            logger.info(
                "Page %d has insufficient text (%d chars), falling back to OCR",
                page_num + 1,
                len(text),
            )
            ocr_text = _ocr_pdf_page(page, page_num, ocr_timeout, warnings)
            if ocr_text:
                page_texts.append(ocr_text)
                used_ocr = True
            else:
                # OCR also failed or returned nothing
                page_texts.append(text)  # Keep whatever little text there was
                if not text:
                    warnings.append(
                        f"Page {page_num + 1}: Could not extract text (direct or OCR)"
                    )

    doc.close()

    # Determine extraction method
    if used_ocr and used_direct:
        method: Literal["direct", "ocr", "mixed"] = "mixed"
    elif used_ocr:
        method = "ocr"
    else:
        method = "direct"

    full_text = "\n\n".join(page_texts)

    return ExtractionResult(
        text=full_text,
        page_count=page_count,
        method=method,
        warnings=warnings,
    )


def _ocr_pdf_page(
    page: fitz.Page,
    page_num: int,
    timeout: int,
    warnings: list[str],
    lang: str = "tha+eng",
) -> str:
    """Render a PDF page to image and OCR it.

    Args:
        page: PyMuPDF page object.
        page_num: Zero-based page index (for logging/warnings).
        timeout: OCR timeout in seconds.
        warnings: List to append warnings to.
        lang: Tesseract language string.

    Returns:
        OCR-extracted text, or empty string on failure.
    """
    tmp_path: str | None = None
    try:
        # Render page to a high-resolution PNG image
        # 300 DPI is standard for OCR (default PyMuPDF is 72 DPI, so matrix 300/72 ≈ 4.17)
        mat = fitz.Matrix(300 / 72, 300 / 72)
        pix = page.get_pixmap(matrix=mat)

        # mkstemp + close before writing: NamedTemporaryFile stays open on Windows
        # and pix.save() then fails with Permission denied.
        fd, tmp_path = tempfile.mkstemp(suffix=".png")
        os.close(fd)
        pix.save(tmp_path)
        return ocr_page(tmp_path, lang=lang, timeout=timeout)

    except subprocess.TimeoutExpired:
        warnings.append(
            f"Page {page_num + 1}: OCR timed out after {timeout}s, returning partial result"
        )
        return ""
    except Exception as e:
        logger.warning("OCR failed for page %d: %s", page_num + 1, str(e))
        warnings.append(f"Page {page_num + 1}: OCR failed ({type(e).__name__}: {e})")
        return ""
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


def ocr_page(image_path: str, lang: str = "tha+eng", timeout: int = DEFAULT_OCR_TIMEOUT) -> str:
    """OCR a single page image using Tesseract.

    Calls Tesseract via subprocess for better timeout control.

    Args:
        image_path: Path to the image file to OCR.
        lang: Tesseract language(s) to use (default: Thai + English).
        timeout: Maximum time in seconds for OCR processing.

    Returns:
        Extracted text from the image.

    Raises:
        subprocess.TimeoutExpired: If OCR exceeds the timeout.
        FileNotFoundError: If Tesseract is not installed/found.
        RuntimeError: If Tesseract exits with a non-zero status.
    """
    cmd = [
        "tesseract",
        image_path,
        "stdout",  # Output to stdout instead of file
        "-l", lang,
        "--psm", "1",  # Automatic page segmentation with OSD
    ]

    try:
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=timeout,
            check=False,
        )
    except FileNotFoundError:
        raise FileNotFoundError(
            "Tesseract OCR is not installed or not found in PATH. "
            "Install with: apt-get install tesseract-ocr tesseract-ocr-tha"
        )

    if result.returncode != 0:
        stderr_msg = result.stderr.strip() if result.stderr else "Unknown error"
        raise RuntimeError(f"Tesseract OCR failed (exit code {result.returncode}): {stderr_msg}")

    return result.stdout.strip()


def extract_docx(file_path: str) -> ExtractionResult:
    """Extract text from a DOCX file preserving document structure.

    Extracts paragraphs with heading levels and table content.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        ExtractionResult with structured text content.
    """
    warnings: list[str] = []
    sections: list[str] = []

    try:
        doc = Document(file_path)
    except Exception as e:
        raise ValueError(f"Failed to open DOCX file: {e}") from e

    # Extract paragraphs, preserving heading structure
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = paragraph.style.name if paragraph.style else ""

        if style_name.startswith("Heading"):
            # Extract heading level from style name (e.g., "Heading 1" -> level 1)
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1

            # Format heading with level indicator
            prefix = "#" * level
            sections.append(f"{prefix} {text}")
        else:
            sections.append(text)

    # Extract table content
    for table_idx, table in enumerate(doc.tables):
        table_lines: list[str] = []
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            # Remove duplicate cell text (merged cells produce duplicates)
            unique_cells = _deduplicate_adjacent(cells)
            if any(unique_cells):  # Skip empty rows
                table_lines.append(" | ".join(unique_cells))

        if table_lines:
            sections.append(f"\n[Table {table_idx + 1}]")
            sections.extend(table_lines)

    full_text = "\n".join(sections)

    # Estimate page count from content (DOCX doesn't have fixed pages)
    # Approximate: ~3000 chars per page for Thai text
    estimated_pages = max(1, len(full_text) // 3000 + (1 if len(full_text) % 3000 else 0))

    return ExtractionResult(
        text=full_text,
        page_count=estimated_pages,
        method="direct",
        warnings=warnings,
    )


def _deduplicate_adjacent(items: list[str]) -> list[str]:
    """Remove adjacent duplicate items (common in merged DOCX table cells).

    Args:
        items: List of cell text strings.

    Returns:
        List with adjacent duplicates removed.
    """
    if not items:
        return []

    result = [items[0]]
    for item in items[1:]:
        if item != result[-1]:
            result.append(item)
    return result


def extract_text_file(file_path: str, encoding: str = "utf-8") -> ExtractionResult:
    """Extract text from a plain text file.

    Args:
        file_path: Path to the text file.
        encoding: File encoding (default UTF-8).

    Returns:
        ExtractionResult with file text content.

    Raises:
        UnicodeDecodeError: If the file cannot be decoded with the given encoding.
    """
    warnings: list[str] = []

    try:
        text = Path(file_path).read_text(encoding=encoding)
    except UnicodeDecodeError:
        # Try Thai encoding fallback (TIS-620 / Windows-874)
        try:
            text = Path(file_path).read_text(encoding="cp874")
            warnings.append(
                "File was not UTF-8 encoded; decoded using Windows-874 (Thai) encoding"
            )
        except UnicodeDecodeError as e:
            raise UnicodeDecodeError(
                e.encoding,
                e.object,
                e.start,
                e.end,
                f"Could not decode file with UTF-8 or Windows-874: {e.reason}",
            )

    return ExtractionResult(
        text=text.strip(),
        page_count=1,
        method="direct",
        warnings=warnings,
    )


def _is_kb_json(path: Path, mime_type: str) -> bool:
    """True when the file should go through a knowledge-base JSON extractor."""
    name = path.name
    if name.endswith(("_tor_extract.json", "_combined.json")):
        return True
    if path.suffix.lower() == ".json" and "04-decision-rules" in path.parts:
        return True
    return mime_type in ("application/json", "text/json")


def _extract_kb_json(file_path: str) -> ExtractionResult:
    """Route combined packs, decision-rules, and TOR extracts."""
    path = Path(file_path)
    name = path.name
    if name.endswith("_combined.json"):
        return extract_combined_kb_json(file_path)
    if "04-decision-rules" in path.parts:
        return flatten_decision_rules_json(file_path)
    return extract_tor_extract_json(file_path)


def _is_raw_json_token(text: str) -> bool:
    """True for identifier-like strings (keys, codes) that should not be embedded."""
    if not text:
        return True
    allowed = set(
        "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789._-"
    )
    return all(char in allowed for char in text)


def _append_thai_text(value: object, parts: list[str]) -> None:
    """Collect Thai (and other prose) strings from nested JSON, skipping raw keys."""
    if isinstance(value, str):
        text = value.strip()
        if text and not _is_raw_json_token(text):
            parts.append(text)
        return
    if isinstance(value, list):
        for item in value:
            _append_thai_text(item, parts)
        return
    if not isinstance(value, dict):
        return
    for key, item in value.items():
        if isinstance(key, str):
            key_text = key.strip()
            if key_text and not _is_raw_json_token(key_text):
                parts.append(key_text)
        _append_thai_text(item, parts)


def extract_combined_kb_json(file_path: str) -> ExtractionResult:
    """Extract RAG text from a knowledge-base `*_combined.json` file.

    These files store Thai excerpts under `sections[].content`.
    Empty packs return empty text so ingestion can fail.
    """
    payload = json.loads(Path(file_path).read_text(encoding="utf-8"))
    contents: list[str] = []
    if isinstance(payload, dict):
        name = payload.get("name")
        if isinstance(name, str) and name.strip():
            contents.append(name.strip())
        sections = payload.get("sections")
        if isinstance(sections, list):
            for section in sections:
                if not isinstance(section, dict):
                    continue
                content = section.get("content") or ""
                if str(content).strip():
                    contents.append(str(content).strip())
    text = "\n\n".join(contents).strip()
    return ExtractionResult(
        text=text,
        page_count=1,
        method="direct",
        warnings=[] if text else ["JSON extract has no sections content"],
    )


def flatten_decision_rules_json(file_path: str) -> ExtractionResult:
    """Flatten nested decision-rules JSON into Thai prose for RAG.

    Nested `rules` objects keep condition/result/legal_basis text and drop
    raw English keys such as `step_1_primary_method` or `R1`.
    Empty payloads return empty text so ingestion can fail.
    """
    payload = json.loads(Path(file_path).read_text(encoding="utf-8"))
    parts: list[str] = []
    _append_thai_text(payload, parts)
    text = "\n\n".join(parts).strip()
    return ExtractionResult(
        text=text,
        page_count=1,
        method="direct",
        warnings=[] if text else ["JSON extract has no rule content"],
    )


def extract_tor_extract_json(file_path: str) -> ExtractionResult:
    """Extract RAG text from a knowledge-base `*_tor_extract.json` file.

    These files store Thai legal excerpts under `focus_areas.*.content`.
    Empty extracts (no focus-area content) return empty text so ingestion
    can fail instead of embedding raw JSON keys.
    """
    payload = json.loads(Path(file_path).read_text(encoding="utf-8"))
    contents: list[str] = []
    if isinstance(payload, dict):
        focus = payload.get("focus_areas")
        if isinstance(focus, dict):
            for sections in focus.values():
                if not isinstance(sections, list):
                    continue
                for section in sections:
                    if not isinstance(section, dict):
                        continue
                    content = section.get("content") or ""
                    if str(content).strip():
                        contents.append(str(content).strip())
    parts = contents
    source = payload.get("source_file") if isinstance(payload, dict) else None
    if source and contents:
        parts = [str(source), *contents]
    text = "\n\n".join(parts).strip()
    return ExtractionResult(
        text=text,
        page_count=1,
        method="direct",
        warnings=[] if text else ["JSON extract has no focus_areas content"],
    )
